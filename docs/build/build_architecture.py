"""
Builds docs/NeurAI-System-Architecture.docx — the current-state reference.

What this document is: everything the platform IS, right now. Every process,
every route family, every table, every technique, with the file that proves
it. What it deliberately is NOT: a history. No decision log, no rationale for
paths not taken, no record of how anything came to be — those live in
ARCHITECTURE.md and the repo's own log.

Run:  python docs/build/build_architecture.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
DIAGRAMS = ROOT / "docs" / "diagrams" / "png"
SHOTS = ROOT / "docs" / "screenshots"
OUT = ROOT / "docs" / "NeurAI-System-Architecture.docx"

# The platform's own palette, so the document and the product agree.
VIOLET = RGBColor(0x5B, 0x34, 0xC7)
VIOLET_LIGHT = "EFE8FF"
INK = RGBColor(0x16, 0x12, 0x1F)
MUTED = RGBColor(0x6B, 0x64, 0x78)
RULE = "DAD4E4"
WASH = "F7F4FC"
CODE_BG = "F2EFF7"

BODY = "Segoe UI"
MONO = "Consolas"


# ---------------------------------------------------------------- helpers

def shade(cell, hex_fill: str) -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def cell_borders(cell, colour: str = RULE, size: int = 4) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(size))
        e.set(qn("w:color"), colour)
        borders.append(e)
    tcPr.append(borders)


def para_rule(paragraph, colour: str = RULE, size: int = 6) -> None:
    """A hairline under a paragraph — the section divider."""
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), colour)
    borders.append(bottom)
    pPr.append(borders)


def styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.24

    for name, size, colour, before, after in (
        ("Heading 1", 20, VIOLET, 22, 8),
        ("Heading 2", 14.5, INK, 16, 6),
        ("Heading 3", 11.5, VIOLET, 12, 4),
    ):
        st = doc.styles[name]
        st.font.name = BODY
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = colour
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    para_rule(p, RULE, 8)
    return p


def h2(doc, text):
    return doc.add_heading(text, level=2)


def h3(doc, text):
    return doc.add_heading(text, level=3)


def body(doc, text, *, size=10.5, colour=INK, italic=False, space=7):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = colour
    r.italic = italic
    return p


def rich(doc, chunks, *, space=7, size=10.5):
    """A paragraph from (text, style) pairs: '', 'b', 'c' (code), 'm' (muted)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    for text, kind in chunks:
        r = p.add_run(text)
        r.font.size = Pt(size if kind != "c" else size - 0.7)
        if kind == "b":
            r.bold = True
            r.font.color.rgb = INK
        elif kind == "c":
            r.font.name = MONO
            r.font.color.rgb = VIOLET
        elif kind == "m":
            r.font.color.rgb = MUTED
        else:
            r.font.color.rgb = INK
    return p


def bullets(doc, items, *, bold_lead=True):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.7)
        if bold_lead and " — " in item:
            lead, rest = item.split(" — ", 1)
            r = p.add_run(lead)
            r.bold = True
            r.font.size = Pt(10.5)
            r2 = p.add_run(" — " + rest)
            r2.font.size = Pt(10.5)
        else:
            p.add_run(item).font.size = Pt(10.5)


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(9)
    r = p.add_run(text)
    r.font.name = MONO
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x35, 0x2C, 0x46)
    return p


def note(doc, title, text):
    """A pulled-out fact — the thing a reader should carry away."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    shade(cell, VIOLET_LIGHT)
    cell_borders(cell, "D9CEFF", 4)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = VIOLET
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        shade(cell, WASH)
        cell_borders(cell)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = VIOLET
    for row in rows:
        cells = t.add_row().cells
        for i, text in enumerate(row):
            cell = cells[i]
            cell_borders(cell)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            mono = text.startswith("`") and text.endswith("`")
            r = p.add_run(text.strip("`"))
            r.font.size = Pt(8.6)
            if mono:
                r.font.name = MONO
                r.font.color.rgb = VIOLET
            else:
                r.font.color.rgb = INK
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


def figure(doc, filename, caption, width_cm=16.4):
    path = DIAGRAMS / filename
    if not path.exists():
        body(doc, f"[missing figure: {filename}]", colour=MUTED, italic=True)
        return
    doc.add_picture(str(path), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(caption)
    r.font.size = Pt(8.6)
    r.font.color.rgb = MUTED
    r.italic = True


def shot(doc, filename, caption, width_cm=15.4):
    path = SHOTS / filename
    if not path.exists():
        return False
    doc.add_picture(str(path), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(caption)
    r.font.size = Pt(8.6)
    r.font.color.rgb = MUTED
    r.italic = True
    return True


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def footer(doc):
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("NeurAI Platform · System architecture · current state, 26 August 2026")
        r.font.size = Pt(8)
        r.font.color.rgb = MUTED


def build() -> Path:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
    styles(doc)

    import content_architecture as C
    C.cover(doc)
    C.contents(doc)
    C.section_topology(doc)
    C.section_identity(doc)
    C.section_pipeline(doc)
    C.section_permissions(doc)
    C.section_agent(doc)
    C.section_refusals(doc)
    C.section_frontend(doc)
    C.section_api(doc)
    C.section_data(doc)
    C.section_ops(doc)
    C.section_screens(doc)
    C.closing(doc)

    footer(doc)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print("wrote", path)
