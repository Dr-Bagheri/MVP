"""
Builds docs/NeurAI-Platform-Demo-Script.docx — what to say over the deck.

One page per slide: the words, the timing, what to point at, and the
questions that slide tends to provoke with an answer for each.

Run:  python docs/build/build_script.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
RENDER = ROOT / "docs" / "build" / "deck-render"
OUT = ROOT / "docs" / "NeurAI-Platform-Demo-Script.docx"

VIOLET = RGBColor(0x5B, 0x34, 0xC7)
INK = RGBColor(0x16, 0x12, 0x1F)
MUTED = RGBColor(0x6B, 0x64, 0x78)
RULE = "DAD4E4"
WASH = "F7F4FC"
PALE = "EFE8FF"
BODY = "Segoe UI"


def shade(cell, fill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(el)


def borders(cell, colour=RULE):
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), colour)
        b.append(e)
    tcPr.append(b)


def styles(doc):
    n = doc.styles["Normal"]
    n.font.name = BODY
    n.font.size = Pt(11)
    n.font.color.rgb = INK
    n.paragraph_format.space_after = Pt(8)
    n.paragraph_format.line_spacing = 1.3


def slide_page(doc, number, name, timing, say, point, asked, *, first=False):
    if not first:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # header strip: number, name, timing
    t = doc.add_table(rows=1, cols=2)
    left, right = t.rows[0].cells
    for c in (left, right):
        borders(c, "FFFFFF")
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"SLIDE {number}")
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = VIOLET
    p2 = left.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(name)
    r2.bold = True
    r2.font.size = Pt(16)
    r2.font.color.rgb = INK
    p3 = right.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p3.paragraph_format.space_after = Pt(0)
    r3 = p3.add_run(timing)
    r3.font.size = Pt(9.5)
    r3.font.color.rgb = MUTED
    left.width = Cm(12)
    right.width = Cm(4.6)

    # the slide itself
    shot = RENDER / f"Slide{number}.JPG"
    if shot.exists():
        doc.add_picture(str(shot), width=Cm(11.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.paragraphs[-1].paragraph_format.space_before = Pt(8)
        doc.paragraphs[-1].paragraph_format.space_after = Pt(10)

    # what to say
    h = doc.add_paragraph()
    h.paragraph_format.space_after = Pt(3)
    hr = h.add_run("SAY")
    hr.bold = True
    hr.font.size = Pt(9)
    hr.font.color.rgb = VIOLET
    for para in say:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        p.paragraph_format.left_indent = Cm(0.3)
        r = p.add_run(para)
        r.font.size = Pt(11.5)

    # what to do
    if point:
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(4)
        h.paragraph_format.space_after = Pt(3)
        hr = h.add_run("POINT AT")
        hr.bold = True
        hr.font.size = Pt(9)
        hr.font.color.rgb = VIOLET
        for item in point:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.8)
            p.add_run(item).font.size = Pt(10.5)

    # likely questions
    if asked:
        t = doc.add_table(rows=1, cols=1)
        cell = t.rows[0].cells[0]
        shade(cell, WASH)
        borders(cell, "E2DAF2")
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("IF THEY ASK")
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = VIOLET
        for q, a in asked:
            pq = cell.add_paragraph()
            pq.paragraph_format.space_after = Pt(1)
            rq = pq.add_run(q)
            rq.bold = True
            rq.font.size = Pt(10.5)
            pa = cell.add_paragraph()
            pa.paragraph_format.space_after = Pt(7)
            ra = pa.add_run(a)
            ra.font.size = Pt(10.5)
            ra.font.color.rgb = INK


def build() -> Path:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.9)
        section.bottom_margin = Cm(1.9)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
    styles(doc)

    import content_script as C
    C.cover(doc)
    C.how_to_use(doc)
    C.build_pages(doc)

    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("NeurAI Platform · demonstration script · August 2026")
        r.font.size = Pt(8)
        r.font.color.rgb = MUTED

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print("wrote", build())
