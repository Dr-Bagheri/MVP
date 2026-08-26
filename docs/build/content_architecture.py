"""The architecture document's content. Imported by build_architecture.py."""
from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from build_architecture import (  # noqa: F401
    INK, MUTED, VIOLET, OUT, SHOTS,
    body, bullets, code, figure, footer, h1, h2, h3, note, page_break,
    para_rule, rich, shot, styles, table,
)


def cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NeurAI Platform")
    r.font.size = Pt(34)
    r.bold = True
    r.font.color.rgb = VIOLET

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("System Architecture")
    r.font.size = Pt(22)
    r.font.color.rgb = INK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("The platform as it stands — every process, every boundary,\n"
                  "every technique, and the file that proves each one.")
    r.font.size = Pt(11.5)
    r.font.color.rgb = MUTED

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    r = p.add_run("26 August 2026")
    r.font.size = Pt(10)
    r.font.color.rgb = MUTED
    page_break(doc)


def contents(doc):
    h1(doc, "What this document is")
    body(doc,
         "A description of the NeurAI Platform exactly as it exists today. It walks the "
         "system from the browser to the database and out to every provider, explains "
         "each part in full, and names the technique each part uses along with the file "
         "that implements it.")
    body(doc,
         "It is a description, not a history. There is no record here of alternatives "
         "considered or paths not taken; where a choice needs explaining, the explanation "
         "is about how the thing works today and why it holds, not about how it came to be.")

    h2(doc, "How to read it")
    bullets(doc, [
        "Sections 1–3 — the shape of the system: what runs where, how a person becomes "
        "an identity, and how a recording becomes a searchable record.",
        "Sections 4–6 — the walls: permissions, the assistant's limits, and what the "
        "platform refuses to do.",
        "Sections 7–9 — the surfaces: the front end, the API, and the data model.",
        "Section 10 — operations: deployment, environment, and the checks that run "
        "before anything ships.",
    ])

    h2(doc, "The product in one paragraph")
    body(doc,
         "NeurAI is a Persian-first platform for organisations. Its first application, "
         "Echo, turns meetings and calls into transcripts, and transcripts into versioned "
         "summaries with action items, decisions and named speakers. An assistant sits "
         "across the whole platform and can answer questions about any record the person "
         "asking is allowed to see — and, with their explicit confirmation, change one.")

    note(doc, "The single most important property",
         "Nothing in this system reads or writes data without a user identity attached. "
         "The database enforces this itself, so the guarantee holds even if every layer "
         "above it is wrong.")
    page_break(doc)


def section_topology(doc):
    h1(doc, "1 · What runs, and where")
    body(doc,
         "The platform is five processes across three hosts, plus one managed database "
         "and two external providers. Nothing else is running in production.")
    figure(doc, "fig-topology.png",
           "Figure 1 — Every production process, its host, and the protocol between them.")

    h2(doc, "1.1 The edge — Vercel")
    rich(doc, [
        ("The web application is ", ""), ("Next.js 16.3", "b"),
        (" on the App Router, deployed to Vercel and served at ", ""),
        ("neurai.pt", "c"), (". It does two jobs that are worth separating.", ""),
    ])
    bullets(doc, [
        "It renders the interface — React 19, two locales, right-to-left first.",
        "It is the Backend-for-Frontend: 103 route handlers under /api that stand "
        "between the browser and the real API.",
    ])
    body(doc,
         "The second job is the security-relevant one. The browser never holds an access "
         "token and never learns the API's address. It calls our own server; our server "
         "attaches the credential and calls the API. This is covered in full in §2.3.")

    h2(doc, "1.2 The origin — a single Linux host")
    body(doc,
         "Four systemd services run on one Hetzner machine. It has no inbound ports open "
         "to the internet: the API is published through a Cloudflare Tunnel, and "
         "everything else listens on loopback only.")
    table(doc,
          ["Service", "What it is", "Listens on"],
          [
              ["`neurai-api`", "The Fastify API. Every product rule lives here.", "`:8080` — tunnelled"],
              ["`neurai-worker`", "Consumes the queues: transcribe, link speakers, summarise, signals.", "no port"],
              ["`neurai-ml`", "The speech service — voice detection, diarization, embeddings.", "`:7801` — loopback"],
              ["`neurai-tts`", "Persian speech synthesis.", "`:5001` — loopback"],
              ["`neurai-purge`", "A timer, not a service. Runs at 03:30, deletes what has expired, exits.", "no port"],
          ],
          widths=[3.6, 9.4, 3.6])
    body(doc,
         "The purge process is separate on purpose and is discussed in §6.3: it is the "
         "only thing in the entire system that can delete a row.")

    h2(doc, "1.3 The managed layer — Supabase")
    body(doc,
         "PostgreSQL, authentication and object storage come from Supabase. Three things "
         "about how they are used matter more than the fact of using them.")
    bullets(doc, [
        "The database is the wall — not a store. Row-level security policies decide what "
        "each identity may see, and they are the enforcement, not a convenience.",
        "PostgREST is not used. The schema is revoked from Supabase's anon and "
        "authenticated roles entirely, so the generated API cannot see that it exists.",
        "Storage has no public bucket policy. Audio is reachable only through a signed, "
        "expiring URL minted by the API.",
    ])

    h2(doc, "1.4 The providers")
    body(doc,
         "Two external services are called, both from the origin host and never from the "
         "browser. Soniox transcribes Persian speech. OpenRouter provides the language "
         "models used for summaries and the assistant. Neither key exists anywhere a "
         "browser can reach.")
    page_break(doc)


def section_identity(doc):
    h1(doc, "2 · Identity — how a person becomes a database row")
    body(doc,
         "Every request in this system carries an identity, and every layer re-derives "
         "what that identity may do rather than trusting what it was told. This section "
         "walks the path once, then explains each technique in it.")
    figure(doc, "fig-auth.png",
           "Figure 2 — Sign-in to a database identity, and the four checks along the way.")

    h2(doc, "2.1 JSON Web Tokens, and how they are verified")
    body(doc,
         "A JWT is a signed statement of who someone is. It carries claims — a subject, "
         "an issuer, an expiry — and a signature. Anyone can read the claims; only the "
         "holder of the signing key can produce a valid signature. Verification is "
         "therefore not a lookup but a cryptographic operation.")
    rich(doc, [
        ("The verifier is hand-written, in ", ""), ("core/src/api/jwt.ts", "c"),
        (". It has ", ""), ("two branches, each pinned to exactly one algorithm", "b"),
        (", and a fall-through that refuses anything that is not exactly one of them. "
         "That structure is the defence: an attacker cannot present a token whose header "
         "claims a weaker algorithm and have the verifier obligingly select it, because "
         "each branch verifies with a key type the other cannot use.", ""),
    ])

    h3(doc, "The ES256 path — the one in use")
    body(doc,
         "Tokens are signed with ECDSA over the P-256 curve. The public keys are published "
         "by the authentication service as a JWKS — a JSON Web Key Set — and each token's "
         "header names which key signed it, by its key id.")
    bullets(doc, [
        "A token with no key id is refused by name, not by falling through to a default.",
        "The key set is cached with a ten-minute lifetime and refetched on a key-id miss, "
        "so a key rotation is survived without a restart and without a thundering herd — "
        "a single in-flight promise is shared by every waiting request.",
        "Only P-256 keys are imported from the set. Importing an RSA key would give an "
        "RS256 claim somewhere to land.",
    ])
    note(doc, "One line that decides whether anything works at all",
         "The signature is verified with dsaEncoding: \"ieee-p1363\". A JWS ECDSA "
         "signature is the raw 64-byte r‖s pair; Node's default is DER. Without that "
         "option every valid token is rejected as a bad signature — the failure looks "
         "like a key problem and is not one.")

    h3(doc, "What is checked, in order")
    bullets(doc, [
        "Expiry, with a 30-second leeway for clock drift.",
        "Not-before, if present.",
        "Issuer — pinned, so a token from another project is not a token here.",
        "Audience — pinned. This is what stops a project's public anonymous key being "
        "presented as a person.",
        "Subject — must be a non-empty string.",
    ])

    h2(doc, "2.2 The token proves who, never what")
    rich(doc, [
        ("After the signature and claims pass, the token's job is finished. Membership, "
         "role and status are re-read from the database on every request — ", ""),
        ("core/src/db/actor.ts", "c"),
        (". A role in a token would be a role a user could still hold an hour after it "
         "was taken away from them.", ""),
    ])

    h2(doc, "2.3 The Backend-for-Frontend")
    body(doc,
         "A BFF is a server that belongs to one front end and exists to talk to the real "
         "back end on its behalf. It is not a general gateway and it holds no rules of "
         "its own.")
    rich(doc, [
        ("Ours is 103 route handlers under ", ""), ("web/src/app/api", "c"),
        (", all routed through one function — ", ""), ("web/src/server/core.ts", "c"),
        (". That function reads the session cookie, attaches the bearer token, and calls "
         "the API. It makes ", ""), ("no authorization decision", "b"),
        (". It carries identity; it does not interpret it.", ""),
    ])
    body(doc, "Three things this buys:")
    bullets(doc, [
        "The browser never holds a credential — an injected script cannot read a token "
        "that is not there.",
        "The API's address is a server-side environment variable, never shipped to a "
        "client bundle.",
        "There is exactly one place where an upstream error becomes a user-facing one, "
        "so the taxonomy cannot drift per screen.",
    ])

    h2(doc, "2.4 The session cookie")
    rich(doc, [
        ("One cookie, ", ""), ("echo_session", "c"), (": ", ""),
        ("httpOnly", "b"), (" so JavaScript cannot read it, ", ""),
        ("sameSite=lax", "b"), (" so it is not sent on cross-site form posts, ", ""),
        ("secure", "b"), (" in production, thirty-day lifetime.", ""),
    ])
    body(doc,
         "The edge middleware reads the cookie's contents rather than its presence. An "
         "unreadable cookie is deleted rather than re-read forever. A token within a "
         "minute of expiry is refreshed at the edge and the renewed cookie is set on the "
         "same response, so the person walks through without noticing. A refusal from the "
         "auth service deletes the cookie; a network failure to that service does not — "
         "a transient outage must not sign everyone out.")

    h2(doc, "2.5 The identity inside the database session")
    body(doc,
         "Having verified who is calling, the API tells the database. It does this with a "
         "transaction-local setting, not a connection-level one, because connections are "
         "pooled and a connection-level identity would leak to whoever borrowed the "
         "connection next.")
    code(doc,
         "set local role echo_app;\n"
         "select set_config('echo.actor_id', $1, true);   -- true = transaction-local")
    rich(doc, [
        ("Every security policy in the schema is written against ", ""),
        ("echo.actor_id()", "c"),
        (". When it is null — which is the state of any connection that did not go "
         "through this door — every policy denies. There is no ambient access.", ""),
    ])
    page_break(doc)


def section_pipeline(doc):
    h1(doc, "3 · From a microphone to a record")
    body(doc,
         "This is the product's core loop: audio arrives, and a searchable, summarised, "
         "attributed record comes out. Every step is a queue message, which is what makes "
         "the loop survivable.")
    figure(doc, "fig-pipeline.png",
           "Figure 3 — The pipeline, and what each step does when it cannot do its job.")

    h2(doc, "3.1 Capture")
    body(doc,
         "Recording happens in the browser, in thirty-minute parts. Parts exist for two "
         "reasons: a long recording that fails should not lose everything, and a hosting "
         "platform's request-size limit is smaller than a full meeting.")
    body(doc,
         "Each part is uploaded directly to object storage using a short-lived signed URL "
         "the API mints, then registered with the API. The audio bytes never pass through "
         "the web server.")

    h2(doc, "3.2 The queue")
    rich(doc, [
        ("Work is handed off through ", ""), ("pgmq", "b"),
        (" — a message queue that lives inside PostgreSQL. Using the database we already "
         "have, rather than a separate broker, means a job and the rows it will write are "
         "in the same system, and there is one thing to operate rather than two.", ""),
    ])
    table(doc,
          ["Queue", "Carries"],
          [
              ["`echo_process_part`", "One part: transcode, detect voice, transcribe, diarize."],
              ["`echo_link_speakers`", "Give every voice a playable snippet; match enrolled voiceprints."],
              ["`echo_summarize`", "Produce a summary version from the transcript."],
              ["`echo_agent_rules`", "Scheduled assistant work — the post-call brief, the weekly digest."],
              ["`echo_deliver_webhook`", "An outbound delivery to a customer's endpoint."],
          ],
          widths=[5.0, 11.6])
    note(doc, "The enqueue contract",
         "Every message carries both the call and the identity it must run as, written "
         "when a genuine caller was present. The worker never performs a privileged read "
         "to work out who a job belongs to — that would be a database access without a "
         "user identity, which is the one thing the system does not do.")

    h2(doc, "3.3 Speech")
    body(doc,
         "The speech service is deliberately productless: it has no database, no identity, "
         "no product credential and no memory. It receives audio, returns a result, and "
         "deletes its workspace. Nothing about a job survives the response.")
    h3(doc, "Voice activity detection")
    body(doc,
         "Silero VAD, an ONNX model, finds the speech in the audio so the silence is never "
         "sent to a paid transcriber. It reduces transcription cost by about fifteen "
         "percent. When the model is unavailable an energy-threshold fallback runs "
         "instead — and the health endpoint names which engine is live rather than "
         "reporting a boolean, because a fallback that is always available makes a "
         "boolean permanently true.")
    h3(doc, "Transcription")
    body(doc,
         "Soniox transcribes Persian at 2.1% word error rate on our reference recording, "
         "with word-level timing. Lanes are tried in a configured order and every attempt "
         "is recorded in the response's provenance.")
    h3(doc, "Who spoke")
    body(doc,
         "Two-channel audio has speakers by construction. Single-channel audio is "
         "diarized with a segmentation model and a speaker-embedding model. Before either "
         "runs, the audio is checked for being dual-mono — one microphone duplicated into "
         "two channels, which is what phone voice memos produce. Treating that as two "
         "speakers transcribes every word twice, invents two people who are one person, "
         "and doubles the bill.")
    body(doc,
         "Speaker labels coming out of the speech service are local: S1, S2. It never "
         "names a person and has never seen the directory.")

    h2(doc, "3.4 Summaries")
    body(doc,
         "A summary is generated by a language model from the transcript, against a "
         "template the person chose. It is versioned: regenerating produces a new version "
         "rather than replacing one, and every version keeps the model, the template and "
         "any extra instruction that produced it. The transcript is the source of truth "
         "and every derived artifact is rebuildable from it.")

    h2(doc, "3.5 What happens when a step cannot do its job")
    body(doc,
         "The pipeline degrades along a stated ladder rather than failing silently or "
         "inventing a value.")
    table(doc,
          ["Situation", "What happens"],
          [
              ["Word timing is unavailable", "Falls back to line timing, then to an anchored speech span. Never to nothing."],
              ["One part fails", "That part becomes a visible gap. The rest of the call completes."],
              ["No model is configured", "The summary is skipped, flagged retryable, and the reason is stored."],
              ["The owner cannot be resolved", "Nothing is written at all. No identity, no write."],
              ["A message is delivered twice", "The second delivery changes nothing — every step is idempotent."],
          ],
          widths=[5.2, 11.4])
    note(doc, "The rule underneath all of them",
         "Whatever is forfeited is said out loud. What was inferred may degrade; what the "
         "user told us never silently changes. A number a person would act on is never "
         "fabricated to fill a gap — an unknown renders as an unknown, in words.")
    page_break(doc)


def section_permissions(doc):
    h1(doc, "4 · Permissions")
    body(doc,
         "Four layers, each of which must independently permit an operation. Only the "
         "bottom one is a wall; the others are conveniences that make the wall pleasant "
         "to live behind.")
    figure(doc, "fig-permissions.png",
           "Figure 4 — The four layers, the database roles, and the named doors.")

    h2(doc, "4.1 Row-level security")
    body(doc,
         "Row-level security is a PostgreSQL feature that attaches a predicate to a table. "
         "Every query against that table — from anywhere, by any code path — silently "
         "gains the predicate. A row the current identity may not see is not filtered out "
         "of the result; it was never in the result.")
    body(doc,
         "All 32 tables in the schema have it enabled, with 86 policies between them. The "
         "policies are written against the identity in the session, which is why §2.5's "
         "transaction-local setting is load-bearing.")
    note(doc, "A refusal and an absence answer the same way",
         "A row you may not see returns 404, not 403. Distinguishing them would let "
         "someone enumerate what exists by watching which error they get — existence is "
         "itself information.")

    h2(doc, "4.2 Role grants")
    body(doc,
         "Below the policies is a blunter layer: the database role a connection uses "
         "decides which tables and columns are reachable at all. A role with no grant "
         "cannot be talked into an operation by any query, however it is phrased.")
    table(doc,
          ["Role", "What it may do"],
          [
              ["`echo_app`", "The product. Reads and writes as the signed-in person — and holds no DELETE anywhere. Deletion in this product is soft; the application cannot express a physical one."],
              ["`echo_agent`", "The assistant's tool calls. A strictly narrower set than the product's, with no grant at all on the calls table and no DELETE anywhere."],
              ["`echo_purge`", "The 30-day hard purge, and nothing else. The only role holding DELETE, and its policies show it only rows whose window has expired."],
              ["`echo_vendor`", "Platform operations. Holds no table privileges at all — only the right to execute three specific functions."],
          ],
          widths=[3.2, 13.4])
    body(doc,
         "The grants migration opens by revoking everything from the public role and from "
         "Supabase's anonymous and authenticated roles. Nothing is granted by default, so "
         "a table added later without a matching grant fails closed rather than open.")

    h2(doc, "4.3 Named doors")
    body(doc,
         "Some operations legitimately need to do slightly more than the caller's own "
         "permissions allow — deleting a call, merging two people, accepting a new "
         "organisation. These do not get a widened policy. They get a function that runs "
         "with the definer's rights, does exactly one thing, checks the caller itself, "
         "and can be listed.")
    rich(doc, [
        ("There are 58 such functions. The distinction that makes them safe: a widened "
         "policy applies to ", ""), ("every query forever", "b"),
        (", including ones nobody has written yet. A door applies to ", ""),
        ("one operation", "b"),
        (", states its own conditions, and appears in a list you can read.", ""),
    ])
    table(doc,
          ["Door", "What it does"],
          [
              ["`soft_delete_call` · `restore_call`", "Mark a call deleted, or bring it back. Restore is admin-only and raises otherwise. Direct writes to the deleted flag are refused for every application role, including admins."],
              ["`merge_person` · `delete_person`", "Fold a duplicate into the person who stays, or remove one, moving voices and links."],
              ["`register_account` · `redeem_invitation`", "The only ways to become a member. Redemption requires the address to match — a forwarded link must not become a bearer token."],
              ["`tombstone_user`", "Empty a person and replace their address. Owner-only; an admin gets the same refusal a member does."],
              ["`vendor_accept_org` · `vendor_set_org_status`", "Platform-level acceptance and suspension, executable only by the vendor role."],
              ["`platform_purge_org` · `platform_purge_user`", "Immediate erasure, storage objects first. Every platform operation requires a written reason."],
          ],
          widths=[5.6, 11.0])
    note(doc, "Structure over predicate",
         "A constraint makes the wrong state unrepresentable. A predicate only makes it "
         "unlikely. Where the two were both available, this schema chose the constraint.")
    page_break(doc)


def section_agent(doc):
    h1(doc, "5 · The assistant")
    body(doc,
         "An agent that can read an organisation's meetings is a serious thing to build. "
         "The design question is not what it can do but what it cannot, and every answer "
         "below is enforced somewhere other than the prompt.")
    figure(doc, "fig-agent.png",
           "Figure 5 — A question, the tools, and the human confirmation a change requires.")

    h2(doc, "5.1 It borrows the caller's authority, never more")
    body(doc,
         "Every tool the model can call is constructed closed over the identity of the "
         "person who asked. The model never receives an unwrapped tool. A tool reaches "
         "exactly what its caller could reach by hand, because it calls the same code the "
         "REST API calls, under the same wall.")
    body(doc,
         "There are two independent layers. The tool checks its own scope and returns a "
         "refusal as an ordinary result — not an exception, because a refusal is a normal "
         "thing for a model to encounter and should read as one. Separately, a central "
         "policy vetoes calls the active skill did not declare, calls above the caller's "
         "role, and calls past the run's budget.")
    note(doc, "Every attempt counts",
         "The tool-call budget counts attempts, not successes. A blocked call still spends "
         "budget, so a model cannot probe the wall for free.")

    h2(doc, "5.2 Write tools do not write")
    body(doc,
         "Three tools can change a record: correcting a transcript line, editing the "
         "speaker roster, replacing a summary. None of them writes anything. Each "
         "validates the change and returns a proposal describing it, with the before and "
         "the after. A test asserts that a write tool issues no insert and no update.")
    body(doc,
         "The proposal renders as a card in the conversation that produced it. This "
         "placement is deliberate: a proposal outside its own conversation has lost the "
         "sentence that made it approvable. There is no pending-approvals inbox anywhere "
         "in the product.")

    h2(doc, "5.3 What a confirmation guarantees")
    bullets(doc, [
        "Decision first, write second — the human's answer is recorded before anything is "
        "applied, so a replayed confirmation is refused before it can act.",
        "The confirmation carries only the run's identifier. The server re-reads the "
        "proposal from its own record rather than trusting a body that could differ from "
        "what the person saw.",
        "One proposal, one outcome — a second confirmation collides with a primary key. "
        "The refusal is structural, not a check that could be forgotten.",
        "Approval widens content, not power. The confirmed write still runs as the agent "
        "role under the same grants.",
        "A rejection is recorded too. A 'no' is a human decision and belongs in the record.",
    ])
    note(doc, "The agent cannot read the answer",
         "It has no grant on the decisions table. An agent that can read a human's "
         "approvals has been handed a way to turn them into a prompt.")

    h2(doc, "5.4 The conversation is a record")
    body(doc,
         "The person's turn is written before the model runs, so a failed answer leaves "
         "the question standing rather than erasing it. An answer that was cut off is "
         "marked as cut off — an unmarked partial answer is, a week later, "
         "indistinguishable from a complete one. Tool calls appear in the thread as names "
         "only; their arguments quote transcripts, so they live on the narrower audit "
         "surface instead.")
    page_break(doc)


def section_refusals(doc):
    h1(doc, "6 · What the platform refuses to do")
    body(doc,
         "Several behaviours are absent by design. They are listed together because their "
         "absence is a feature, and because an absence is easy to mistake for something "
         "not yet built.")

    h2(doc, "6.1 It does not fabricate a number")
    body(doc,
         "Three states are kept distinct everywhere: not fetched yet, could not be read, "
         "and genuinely empty. Two of those are 'we do not know' and only the third is a "
         "fact about the organisation. A dashboard tile that cannot read its data says so; "
         "it does not show a zero.")

    h2(doc, "6.2 It does not put content in logs")
    body(doc,
         "Transcripts, summaries, message text and speech are never logged. Database "
         "errors are logged by structured field — code, constraint, table, column — never "
         "by message or detail, because the detail field quotes the offending row. The "
         "error reporter is built to carry no message at all for database errors, and to "
         "strip every quoted string from the others.")

    h2(doc, "6.3 Deletion is soft, and erasure is separate")
    body(doc,
         "Deleting a record marks it and sets a purge date. It disappears from the "
         "product immediately and remains restorable by an admin for the window. When the "
         "window expires, a separate process — the only one in the system holding DELETE "
         "— removes it permanently.")
    note(doc, "Objects before rows",
         "The purge deletes storage objects first, then rows. The row is the only pointer "
         "to the object, so a crash in the other order would leave an orphaned file with "
         "nothing left to say it should be removed.")

    h2(doc, "6.4 It does not fetch arbitrary addresses")
    body(doc,
         "Outbound webhooks are a request the customer controls, which makes them a "
         "server-side request forgery risk. The guard runs at the moment the socket opens, "
         "on the address actually being connected to — the only point at which DNS "
         "rebinding cannot walk past a check. Redirects are followed manually and each "
         "hop is re-checked.")
    body(doc,
         "A blocked address is non-retryable. Retrying a block would turn one refusal into "
         "a slow port scan, and the distinction between 'blocked' and 'transport failed' "
         "is the security-relevant part.")

    h2(doc, "6.5 Webhook deliveries are signed, and cannot be replayed forever")
    body(doc,
         "The signature covers a timestamp joined to the body, with a five-minute "
         "tolerance, so a captured delivery is not replayable indefinitely. The secret is "
         "shown once at creation and stored only as a derived hash. There is no reveal "
         "affordance anywhere in the product.")

    h2(doc, "6.6 The assistant's model list excludes one vendor")
    body(doc,
         "A product rule, applied before the organisation's own allow-list so that no "
         "later filter can undo it, with tests that fail if the excluded models reappear.")
    page_break(doc)


def section_frontend(doc):
    h1(doc, "7 · The front end")
    figure(doc, "fig-frontend.png",
           "Figure 6 — Shell, routes, tokens, and what Persian-first means in practice.")

    h2(doc, "7.1 The shell")
    body(doc,
         "One shell wraps every surface: an icon rail of four primary destinations, a "
         "section menu belonging to the current app, and the content with a breadcrumb "
         "trail above it. The assistant is present on every screen as a floating "
         "presence rather than a destination.")
    body(doc,
         "The breadcrumb trail is built from a declared table of routes, not from "
         "splitting the URL. A path-derived trail would teach an information architecture "
         "that the navigation contradicts.")

    h2(doc, "7.2 Surfaces")
    table(doc,
          ["Route", "What it is"],
          [
              ["`/{locale}`", "The dashboard — the landing page. A widget board the person arranges."],
              ["`/{locale}/assistant`", "The assistant: greeting, prompt, conversation history."],
              ["`/{locale}/echo/…`", "Echo: new meeting, records, archive, speakers, search."],
              ["`/{locale}/calls/{id}`", "One record: summary document, transcript, notes, related records."],
              ["`/{locale}/management/…`", "Users, models, skills, server health, connectors."],
              ["`/{locale}/settings/…`", "Configuration, connections, compliance, audit."],
              ["`/{locale}/platform`", "The platform control plane — organisations, users, audit."],
          ],
          widths=[4.6, 12.0])

    h2(doc, "7.3 The design system")
    body(doc,
         "Every colour on screen is a CSS custom property. A component names a role — "
         "surface, foreground, accent, danger — and never a value, which is what lets both "
         "themes resolve as complete sets rather than as an inversion of each other.")
    rich(doc, [
        ("Contrast is not left to judgement. A script checks every foreground/background "
         "pair against its floor, compositing translucent layers rather than comparing hex "
         "values, and ", ""), ("exits non-zero", "b"),
        (". It runs as part of the test suite, so a token change that breaks a pair fails "
         "the build instead of reaching a screen.", ""),
    ])

    h3(doc, "Shared controls")
    body(doc,
         "The table, the menu, the dialog and the dropdown are single shared components. "
         "A screen cannot invent its own dialect of a control. Two rules the shared "
         "components enforce for the whole product: every action lives in one place — "
         "right-click a row — and destructive actions sort to the bottom of every menu, "
         "grouped and separated, so a mis-click cannot land on one in passing.")

    h2(doc, "7.4 Persian-first, in practice")
    bullets(doc, [
        "Direction is structural — layout uses logical properties throughout, so a "
        "component is written once and works in both directions.",
        "Digits follow the language; months follow the calendar preference. Two axes, "
        "deliberately separate, so a Persian speaker can read Gregorian dates if they "
        "work that way.",
        "Names render as authored. A person's name is never transliterated; a second "
        "Latin name is a separate field, never derived from the first.",
        "Text is normalised at ingest and at query, so a search matches regardless of "
        "which Arabic character variant was typed.",
        "Every user-facing string exists in both locales, asserted by a test that reads "
        "the source for every translation call.",
    ])
    note(doc, "A build check for a class of bug that is invisible on screen",
         "A byte-level sweep over every tracked text file fails the build on a "
         "byte-order mark or a mis-encoded character. Encoding damage frequently "
         "survives visual review — a mangled non-breaking space renders as a space — and "
         "the sweep reads bytes rather than looking at text.")

    h2(doc, "7.5 The dashboard")
    body(doc,
         "The landing page is a board of widgets the person arranges. It has three layers, "
         "and the separation is what makes it extensible.")
    bullets(doc, [
        "A registry declares what each widget is: its icon, its colour family, the sizes "
        "it is designed at, which section of the add menu it belongs to. Adding a widget "
        "is an entry there plus a renderer.",
        "A layout model stores where each card sits and validates a stored board against "
        "the registry on read, so a saved layout can never name a widget that no longer "
        "exists.",
        "A grid engine handles collision and reflow. React owns content, the engine owns "
        "geometry, and neither writes the other's half.",
    ])
    body(doc,
         "There are four fixed tile sizes rather than free resizing. A closed set is what "
         "lets each size be designed instead of stretched, and it carries a rule: a bigger "
         "tile shows more information, never the same content scaled up. Each widget "
         "receives its size as a value and renders accordingly.")
    page_break(doc)


def section_api(doc):
    h1(doc, "8 · The API")
    body(doc,
         "One Fastify server, 125 registered routes under /v1. It owns every product "
         "rule; nothing above it re-implements one.")

    h2(doc, "8.1 The guard vocabulary")
    body(doc,
         "Authorization is expressed as a small ladder of named guards, so a route's "
         "requirement is visible at its declaration rather than buried in its body.")
    table(doc,
          ["Guard", "Means"],
          [
              ["`identify`", "The token is verified and membership re-derived. The account may be inactive."],
              ["`requireActive`", "An active member of an active organisation. The default for product routes."],
              ["`requireAdmin`", "Admin or owner."],
              ["`requireOwner`", "The organisation's owner only."],
              ["`requirePlatformRoot`", "Platform operations. Deliberately not built on requireActive, so a root can reactivate their own suspended organisation."],
          ],
          widths=[4.4, 12.2])

    h2(doc, "8.2 Route families")
    table(doc,
          ["Family", "Covers"],
          [
              ["Public", "Health, which sign-in methods exist, registration, invitation redemption."],
              ["Calls", "Create, upload parts, finish, retry, transcript, summaries, speakers, notes, archive, delete."],
              ["Assistant", "Ask (streamed), conversations, messages, feedback, sharing, run traces, proposals."],
              ["Directory", "People, titles, teams, merge, voice enrollment."],
              ["Search", "Full-text across transcripts, Persian-folded, with highlighted extracts."],
              ["Skills & agents", "Resolved skills, authoring, agents, workflows, connectors."],
              ["Admin", "Members, invitations, audit, server health, models, organisation, gateway keys and webhooks."],
              ["Platform", "Organisations, users, audit, purge, root grants. Every operation requires a reason."],
          ],
          widths=[3.8, 12.8])

    h2(doc, "8.3 The error taxonomy")
    body(doc,
         "Errors carry a kind alongside the status, so a client can distinguish cases that "
         "share a code without parsing prose.")
    table(doc,
          ["Status", "Kinds"],
          [
              ["401", "`no_token` · `bad_signature` · `unknown_actor` · `bad_key`"],
              ["403", "`pending` · `suspended` · `forbidden`"],
              ["404", "`not_found` — a row you cannot see and a row that does not exist"],
              ["400", "`invalid`, with a code and parameters so a client never re-implements a server rule"],
              ["409", "`conflict`"],
          ],
          widths=[2.6, 14.0])
    body(doc,
         "The not-found handler emits the same shape as a hidden row, so a missing route "
         "and a forbidden row cannot be told apart by their shape either.")

    h2(doc, "8.4 Streaming")
    body(doc,
         "The assistant answers over server-sent events. The contract has five event "
         "types and one absolute rule: the terminal event is always sent, including on "
         "failure. A client treats a stream that ends without it as a transport failure, "
         "which is only a correct inference if the stream is never dropped silently.")
    page_break(doc)


def section_data(doc):
    h1(doc, "9 · The data model")
    figure(doc, "fig-data.png",
           "Figure 7 — The tables a record touches, and what each is for.")

    h2(doc, "9.1 Shape")
    body(doc,
         "32 tables in one schema, built by 98 numbered migrations. The schema is "
         "hand-written SQL; there is no ORM, and queries are written against the database "
         "directly. Migrations are append-only and checksummed: the runner refuses to "
         "apply a file that has changed since it was recorded.")

    h2(doc, "9.2 The clusters")
    bullets(doc, [
        "Tenancy — the organisation, its members, invitations, and the status history "
        "that lets the product show movement rather than only a current count.",
        "Records — the call, its parts, its transcript segments, and the speakers in it.",
        "Derived — summaries (versioned), the people directory, and the link between a "
        "voice in one call and a person in the directory.",
        "Assistant — conversations, messages, runs, cards, and scheduled rules.",
        "Governance — the audit ledger, human decisions on proposals, the deletion "
        "record, and the platform's own audit.",
        "Integration — API keys, webhooks, deliveries, and connector credentials.",
    ])

    h2(doc, "9.3 Facts worth knowing about the model")
    bullets(doc, [
        "A voiceprint is a vector, never audio — the platform keeps no enrollment "
        "recording. Re-enrolling averages into the existing print rather than replacing "
        "it, so a second sample sharpens the match instead of betting on the newest clip.",
        "A summary version keeps the model, template and instruction that produced it, so "
        "any version can be explained after the fact.",
        "A person's turn in a conversation has no run behind it; the assistant's does. "
        "The link is nullable because that difference is real.",
        "The audit ledger records field names, never values, and is written in the same "
        "transaction as the change it describes — a log in its own transaction can record "
        "an action that failed.",
        "The deletion record survives the purge of what it describes. The content is "
        "erased; the fact that a deletion happened is not content.",
    ])
    page_break(doc)


def section_ops(doc):
    h1(doc, "10 · Operations")

    h2(doc, "10.1 Deploying")
    body(doc,
         "The web application deploys itself: a push to the main branch triggers a Vercel "
         "build. The origin host is updated by extracting a git archive of the release "
         "into the deploy directory, reconciling dependencies, and restarting the "
         "services. Node modules and machine-learning models are not in the archive and "
         "survive redeploys.")
    note(doc, "Schema leads code",
         "A release that adds a migration applies the migration before restarting the "
         "services. The application tolerates a schema that is ahead of it; it does not "
         "tolerate one that is behind.")

    h2(doc, "10.2 Migrations")
    body(doc,
         "Migrations are run by an operator from a workstation using the owner "
         "connection. That credential is deliberately not on the server, so a compromised "
         "application host cannot alter the schema.")

    h2(doc, "10.3 Capability detection")
    body(doc,
         "Deployments and migrations can arrive in either order. The API inspects the "
         "live schema once at boot for the columns and tables that newer features need, "
         "caches the answer, and degrades those features with an honest message rather "
         "than failing. A feature whose column does not exist yet says so; it does not "
         "render an empty state that would read as 'you have none'.")

    h2(doc, "10.4 Secrets")
    body(doc,
         "No credential exists in the repository. Server-side secrets live in root-owned "
         "environment files on the origin host, readable by the service user and nothing "
         "else. The one public value — the browser-safe publishable key — is public by "
         "design and ships in every client bundle of every application built this way.")

    h2(doc, "10.5 What runs before anything ships")
    table(doc,
          ["Package", "The suite"],
          [
              ["core", "Unit and integration tests, including a boot test that starts the real runtime and requires it to answer one request."],
              ["web", "Component tests, the contrast verifier, the byte-level encoding sweep, and a real production build."],
              ["ml", "Model and pipeline tests, with positive-detection assertions — a model wired up wrongly usually passes a negative test."],
              ["db", "A SQL suite that asserts the wall: every table has row-level security, only the purge role holds DELETE, and each door refuses the callers it should."],
          ],
          widths=[2.4, 14.2])
    body(doc,
         "Opt-in lanes exist alongside these for tests that spend real money or need real "
         "network — the live pipeline, the proposal loop, webhook delivery, the purge. "
         "They are deliberately outside the default suite and are run at acceptance.")

    h2(doc, "10.6 Observability")
    body(doc,
         "Errors are reported to a Sentry-protocol endpoint by a zero-dependency "
         "reporter that is built to carry less rather than filtered to carry less: "
         "database error messages never travel, other messages travel truncated with "
         "quoted strings removed, and no request bodies, headers or identifiers beyond a "
         "uuid are sent. With no endpoint configured it is dark.")
    page_break(doc)


def section_screens(doc):
    h1(doc, "11 · The product, on screen")
    body(doc,
         "The surfaces described above, as they render. Both themes and both writing "
         "directions are shown where they differ meaningfully.")
    pairs = [
        ("hub-fa.png", "The assistant hub in Persian — right-to-left, with the presence orb and the prompt."),
        ("hub-en.png", "The same surface in English. The layout is mirrored by direction, not rebuilt."),
        ("new-meeting.png", "Starting a meeting: device check, template choice, and the recorder."),
        ("new-meeting-recording.png", "Recording in progress. The take continues across a navigation."),
        ("topbar-recording.png", "The recorder docks into the top bar, so a recording is never lost by moving."),
        ("record-page.png", "One record: the summary document, its versions, and the transcript beneath."),
        ("records-bulk.png", "The records table with a selection — one width, one menu, one way in."),
        ("speakers-enrollment.png", "The speakers directory, with scripted voice enrollment."),
        ("management-users.png", "Member management: the pending queue first, server-side search and filters."),
        ("management-server.png", "Server health. Every metric carries when it was measured."),
        ("platform-console.png", "The platform control plane — organisations, users, and the audit trail."),
        ("sign-in-dark.png", "Sign-in, dark."),
        ("sign-in-light.png", "Sign-in, light. Both themes are designed, not inverted."),
    ]
    for name, caption in pairs:
        shot(doc, name, caption)


def closing(doc):
    h1(doc, "Appendix · Where things live")
    body(doc,
         "The file that implements each technique described in this document.")
    table(doc,
          ["Technique", "File"],
          [
              ["Token verification", "`core/src/api/jwt.ts`"],
              ["Identity resolution and guards", "`core/src/api/auth.ts` · `core/src/db/actor.ts`"],
              ["Transaction-local identity", "`core/src/db/identity.ts`"],
              ["The BFF hop", "`web/src/server/core.ts`"],
              ["Session cookie and refresh", "`web/src/server/session.ts` · `web/src/middleware.ts`"],
              ["Row-level security", "`db/migrations/0013_rls_policies.sql`"],
              ["Role grants", "`db/migrations/0012_roles.sql` · `0014_grants.sql`"],
              ["Named doors", "`db/migrations/0032` · `0069` · `0076` · `0092` · `0095` · `0096`"],
              ["Queueing", "`core/src/worker/queue.ts` · `db/migrations/0017_queues.sql`"],
              ["The pipeline steps", "`core/src/worker/steps.ts` · `call-steps.ts`"],
              ["The agent runtime", "`core/src/agent/runtime.ts` · `policy.ts` · `tools.ts`"],
              ["Proposals and approvals", "`core/src/agent/proposals.ts` · `db/migrations/0029`"],
              ["Speech", "`ml/src/pipeline.ts` · `ml/CONTRACT.md`"],
              ["Storage signing", "`core/src/storage/signer.ts`"],
              ["Address guard", "`core/src/net/address-guard.ts`"],
              ["Webhook signing", "`core/src/worker/webhook-signing.ts`"],
              ["Purge", "`core/src/purge/purge.ts`"],
              ["Audit", "`core/src/api/admin-actions.ts` · `core/src/api/audit.ts`"],
              ["Design tokens and contrast", "`web/src/app/globals.css` · `design-system/neurai-platform/verify-pairs.mjs`"],
              ["Dashboard registry and board", "`web/src/lib/widgetRegistry.tsx` · `web/src/components/platform/dashboard/WidgetBoard.tsx`"],
          ],
          widths=[5.6, 11.0])
