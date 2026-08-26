"""The speaker script's content."""
from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from build_script import INK, MUTED, VIOLET, slide_page  # noqa: F401


def cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NeurAI Platform")
    r.font.size = Pt(30)
    r.bold = True
    r.font.color.rgb = VIOLET

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Demonstration — speaker script")
    r.font.size = Pt(18)
    r.font.color.rgb = INK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run("Thirteen slides. Twelve to fifteen minutes at a comfortable pace,\n"
                  "leaving room for questions as they come rather than at the end.")
    r.font.size = Pt(11.5)
    r.font.color.rgb = MUTED


def how_to_use(doc):
    doc.add_paragraph().add_run().add_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    r = p.add_run("How to use this")
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = VIOLET

    for text in [
        "The words under SAY are written to be spoken, not read. Say them in your own "
        "phrasing — what matters is the order of the ideas and the two or three "
        "sentences marked as the ones to land exactly.",
        "POINT AT tells you what on the slide to draw attention to, and when. A slide "
        "where you say everything and point at nothing is a slide the audience reads "
        "instead of listening to.",
        "IF THEY ASK covers the questions each slide reliably provokes. Answer them "
        "when they come — an interruption you can answer well is worth more than the "
        "rest of your script.",
        "One rule for the whole talk: never claim more than is built. Where something "
        "is planned rather than shipped, this script says so, and so should you. The "
        "credibility of everything else depends on it.",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(9)
        p.add_run(text).font.size = Pt(11)


SLIDES = [
    (1, "Opening", "45 seconds",
     ["Good morning. I want to show you something we have built and something we are "
      "building toward, and I will try to be precise about which is which.",
      "The line on the screen is the whole idea. The meeting is the record. Not the "
      "notes somebody took in the meeting — the meeting itself, kept in a form you "
      "can search, act on, and point to later.",
      "This is NeurAI. It is a platform, and Echo is the first application inside it."],
     ["The line — read it once, out loud, then pause before continuing."],
     [("Is this an Iranian product?",
       "It is built Persian-first, which is different from being translated into "
       "Persian. Right-to-left layout, Persian digits, Jalali dates and Persian text "
       "normalisation are structural, not a language pack. It works fully in English "
       "too.")]),

    (2, "The problem", "90 seconds",
     ["Every organisation has the same three questions after a meeting, and no reliable "
      "way to answer any of them.",
      "What was decided — nobody wrote it down, and two people remember it differently. "
      "Who committed to what — it was obvious in the room and it is not obvious on "
      "Thursday. Has this come up before — probably, and there is no way to look.",
      "This is the sentence I would like to land: meetings are where organisations "
      "actually decide things, and they are the least recorded part of the business. "
      "Everything else gets a system. This gets somebody's memory."],
     ["The three cards, one at a time, in order.",
      "The bold line at the bottom — say it slowly, it is the argument."],
     [("Don't people already record meetings?",
       "Some do, and then they have an audio file nobody will ever open. A recording "
       "is not a record. The question is whether you can find the one sentence you "
       "need six weeks later.")]),

    (3, "What we built", "75 seconds",
     ["So: two things, and the distinction matters.",
      "On the left, the platform. That is the shell — who you are, which organisation "
      "you belong to, what you are allowed to see, and an assistant that reaches "
      "across all of it. Every application we build sits inside that shell.",
      "On the right, Echo. The first application, and the one you will see today. It "
      "records or ingests a meeting, transcribes it, works out who spoke, and produces "
      "a summary against a template you chose."],
     ["The two panels — the violet one first, then the pale one.",
      "The word 'first' in 'the first application'. That is the roadmap in one word."],
     [("Why build the platform first, when you only have one app?",
       "Because the platform is the expensive part and we would build it anyway. "
       "Permissions, identity, audit and the assistant are the same problem whichever "
       "application asks. Doing it once means the second application inherits it.")]),

    (4, "Why a platform", "90 seconds",
     ["Let me be direct about what is hard here, because it is not the part people "
      "expect.",
      "Transcription is a purchased capability. We buy it, and it is very good — I will "
      "show you the number in a minute. What takes years is everything around it.",
      "Who can see what — an organisation's meetings are its most sensitive data, and "
      "we enforce that in the database itself rather than in the interface. An "
      "assistant with limits — it answers with exactly your access and cannot change "
      "anything without you confirming it. Persian done properly. And a record you can "
      "actually defend, with versions and provenance and an audit trail.",
      "None of that is specific to meetings. All of it is the platform."],
     ["Each row as you say it. Do not read them — you are summarising them.",
      "Pause after 'we enforce that in the database itself'. It is the claim the whole "
      "talk rests on."],
     [("What does 'enforced in the database' actually mean?",
       "Every table has a rule attached to it that says which rows the current user may "
       "see. It applies to every query from every code path, including ones nobody has "
       "written yet. If our own application had a bug, the database would still refuse. "
       "There is a diagram of this later if you want it.")]),

    (5, "How Echo works", "80 seconds",
     ["Four steps, and the person doing them has to remember none of it.",
      "Record — in the browser, or upload a file you already have. It records in "
      "thirty-minute parts, so a two-hour meeting is never one fragile thing.",
      "Transcribe — we strip the silence first, so we are not paying to transcribe "
      "nothing, then Persian speech becomes text with timing on every word.",
      "Attribute — voices are separated and matched against people you have enrolled. "
      "You confirm the rest, once, and next time they are recognised.",
      "Summarise — against the template you chose. Actions, decisions, tags. And it is "
      "versioned: regenerating gives you a new version, it never overwrites the old one.",
      "The line at the bottom is the part that matters operationally. If one part fails, "
      "you get a visible gap in that transcript and the rest of the meeting still "
      "completes. It does not fail the whole thing and it does not quietly pretend the "
      "gap is not there."],
     ["The four numbered cards, in sequence.",
      "The bottom line — this is the one operations people care about."],
     [("What accuracy do you get?",
       "2.1% word error rate on Persian, measured on a real recording after "
       "normalisation. That is two words in a hundred, and in practice most of them are "
       "loanword spellings.")]),

    (6, "The pipeline", "60 seconds",
     ["This is that pipeline as a diagram, and I am showing it for one reason: the "
      "bottom row.",
      "Those four boxes are what happens when a step cannot do its job. Word timing "
      "unavailable — we fall back to line timing, then to a speech span, never to "
      "nothing. One part fails — a visible gap. No model configured — the summary is "
      "skipped and flagged, and it says why. The owner cannot be resolved — nothing is "
      "written at all.",
      "The rule underneath all four is the one I want you to remember: whatever is "
      "forfeited is said out loud. We do not fill a gap with a plausible number."],
     ["The bottom strip — that is the whole slide.",
      "The last line, twice if you have to."],
     [("Why does that matter so much to you?",
       "Because a system that quietly guesses is worse than one that admits it does not "
       "know. If you cannot trust a number on the screen, you have to verify all of "
       "them, and then the system has saved you nothing.")]),

    (7, "The record", "90 seconds",
     ["This is one meeting.",
      "The summary at the top, versioned — regenerate it against a different template "
      "and the old version is still there. The transcript beneath it: click any line "
      "and it plays from that moment, on one timeline across every part of the "
      "recording.",
      "Speakers are people, not labels. Link a voice to somebody once and they are "
      "recognised in future meetings.",
      "And corrections are kept honestly. If somebody edits a line, the record says it "
      "was human-edited and who did it. We do not silently absorb a change into the "
      "transcript as though the machine got it right."],
     ["The summary, then the transcript beneath it.",
      "A speaker name — that is where the directory shows up.",
      "If you are demonstrating live: click a transcript line and let them hear it seek."],
     [("Can we edit the transcript?",
       "Yes, and the edit is marked. The point is not to stop people correcting things "
       "— it is that six months later you can tell what the machine heard from what a "
       "person decided it should say.")]),

    (8, "The assistant", "2 minutes",
     ["This is the part that usually raises the most questions, so let me get ahead of "
      "them.",
      "You can ask it anything about your meetings. 'What did we decide about the "
      "budget?' — it searches the transcripts you are allowed to read and answers with "
      "the meeting it found the answer in.",
      "It cannot see more than you can. Every tool it uses runs under your identity. "
      "There is no privileged path and no admin mode where it sees everything.",
      "And it does not change things. When it wants to — correct a transcript line, fix "
      "a speaker, replace a summary — it describes the change, shows you what it would "
      "look like before and after, and waits. Your answer is recorded either way, "
      "because a 'no' is a decision too.",
      "The line at the bottom is the one to land. This limit is in the database, not in "
      "the prompt. A prompt is a request. A grant is a wall."],
     ["The four cards, left column then right.",
      "The violet strip at the bottom. Say it, then stop talking for a second."],
     [("How do you stop it from making things up?",
       "Two ways. It answers from your actual transcripts and points at the meeting, so "
       "you can check. And for anything that would change data, it cannot act — it can "
       "only propose, and a person approves."),
      ("Which model is it?",
       "Configurable per organisation, from a curated list. One vendor is excluded as a "
       "product rule, applied before your own allow-list so nothing downstream can undo "
       "it.")]),

    (9, "The dashboard", "60 seconds",
     ["The landing page is a board, not a report.",
      "Fifteen widgets, four sizes each, and you drag them where you want them. A tile "
      "you leave at the bottom stays at the bottom — the board does not rearrange "
      "itself under you when you come back.",
      "A bigger tile says more, rather than the same thing larger. Growing a card adds "
      "rows, or range, or a chart.",
      "And the third one is the same principle as the pipeline slide: it does not "
      "invent a number. Not-loaded-yet, could-not-read, and genuinely-empty are three "
      "different states and the board tells you which one it is in."],
     ["The three bold lines, in order.",
      "If demonstrating live: drag a tile, then resize one from its menu."],
     [("Can each person have their own layout?",
       "Yes — the arrangement is per person. Today it is stored in the browser; the "
       "server-side slot for it exists and is the next step for that feature. I would "
       "rather tell you that than imply it is already synced across devices.")]),

    (10, "How it is built", "90 seconds",
     ["For the technical people in the room, this is the whole system.",
      "Five processes. The web application on the edge, the API and workers on a "
      "dedicated host behind a tunnel with no open inbound ports, a speech service, and "
      "a purge process that runs once a night.",
      "Two things worth pointing out. The browser never holds a credential — it talks "
      "to our own server, which attaches the token and calls the API. So there is no "
      "access token in a browser for anything to steal.",
      "And the database is the wall. Every permission is enforced there, which means "
      "the guarantee holds even if every layer above it has a bug."],
     ["The browser box, then follow the arrow to the edge — that is the token point.",
      "The Postgres box, marked 'the wall'.",
      "The purge box, if anyone asks about deletion."],
     [("Where is the data hosted?",
       "Managed Postgres in Frankfurt, with the application on a dedicated host. "
       "Nothing about a customer's meetings leaves that boundary except the audio we "
       "send to the transcription provider and the text we send to the model."),
      ("Is it multi-tenant?",
       "Yes, and the isolation is at the row level in the database rather than by "
       "convention in the application. An organisation's rows are simply not visible to "
       "another organisation's queries.")]),

    (11, "What it refuses to do", "80 seconds",
     ["I want to spend a minute on absences, because they are easy to mistake for "
      "things we have not got to yet.",
      "No content in logs — transcripts, summaries and messages are never written to a "
      "log, and database errors are logged by code rather than by message, because the "
      "message quotes the row.",
      "The application cannot delete a row. It can mark something deleted; one separate "
      "process, running at half past three in the morning, is the only thing in the "
      "system with the right to actually remove it.",
      "No fabricated numbers, which we have covered. No arbitrary outbound requests — "
      "webhook addresses are checked at the moment the connection opens, on the address "
      "actually being reached. And no token in the browser.",
      "Every one of those is a decision that cost us something to keep."],
     ["Three or four of the six — do not read all of them.",
      "The deletion one is usually the one they want to talk about."],
     [("What about GDPR / data deletion requests?",
       "Deleting a record removes it from the product immediately and permanently "
       "erases it when the window expires. There is also an immediate purge at the "
       "platform level for an organisation or a person, which removes stored audio "
       "before the rows — and every platform operation requires a written reason that "
       "is recorded.")]),

    (12, "Where it stands", "60 seconds",
     ["Numbers, briefly, because I would rather be specific than enthusiastic.",
      "2.1% word error rate on Persian, measured rather than quoted. Ninety-eight "
      "database migrations, all hand-written and checksummed. Thirty-two tables, every "
      "one of them walled. And over fourteen hundred automated checks across four "
      "suites, plus a real production build, before anything ships.",
      "It is running in production today. Not a prototype and not a demo environment — "
      "this is the system."],
     ["The four numbers.",
      "The strip at the bottom — 'running in production today'."],
     [("How many customers?",
       "Be honest here — say the real number. The system is live and in use; the "
       "commercial stage is what it is, and claiming otherwise is the fastest way to "
       "lose the room.")]),

    (13, "Where we are heading", "75 seconds",
     ["Three columns, and I have been careful about which is which.",
      "Now — everything you have seen. Record, transcribe, attribute, summarise, search, "
      "and an assistant that answers and proposes. That is built and running.",
      "Next — the assistant stops waiting to be asked. Questions raised in a meeting "
      "that nobody answered. A decision that reverses an earlier one. Preparation before "
      "a meeting rather than a summary after it. Terms you choose to watch for.",
      "Then — more applications. The shell is application-agnostic on purpose. Identity, "
      "permissions and the assistant are already shared, so the second application "
      "inherits the wall instead of rebuilding it."],
     ["The three columns, left to right. Be visibly clearer about the first than the "
      "other two.",
      "'Inherits the wall' — that is why the platform decision pays off."],
     [("When is 'next'?",
       "Give a range you can keep, not a date you cannot. The honest framing is that "
       "these are designed and the data they need already exists — the work is in the "
       "detection being right often enough to be worth trusting.")]),
]


def build_pages(doc):
    for i, (n, name, timing, say, point, asked) in enumerate(SLIDES):
        slide_page(doc, n, name, timing, say, point, asked, first=False)
