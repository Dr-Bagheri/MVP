from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
OUT = ROOT / "deliverables" / "neurai-platform-final"
ASSETS = OUT / "build" / "assets"
SCREEN = ROOT / "docs" / "screenshots"
BRAND = ROOT / "web" / "public" / "brand" / "neurai-mark-light-transparent.png"
HERO = ASSETS / "ai-native-hero.png"

PURPLE = "9B6DFF"
PURPLE_DARK = "251044"
PURPLE_PALE = "F1ECFF"
INDIGO = "0C0922"
INK = "171425"
SLATE = "4C465C"
MUTED = "786F89"
LINE = "DED8E9"
PAPER = "FBFAFD"
WHITE = "FFFFFF"
CORAL = "FF6F61"
CYAN = "63D8FF"
GREEN = "2E9B72"
AMBER = "C47A13"
RED = "B93B53"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key in ["val", "sz", "space", "color"]:
            if key in kwargs[edge]:
                element.set(qn("w:" + key), str(kwargs[edge][key]))


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def keep_with_next(paragraph, value=True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn("w:keepNext"))
    if node is None:
        node = OxmlElement("w:keepNext")
        p_pr.append(node)
    if not value:
        node.set(qn("w:val"), "0")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell, width) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.twips)))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    e = OxmlElement("w:tblHeader")
    e.set(qn("w:val"), "true")
    tr_pr.append(e)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_rtl(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def add_hyperlink(paragraph, text: str, url: str, color=PURPLE, underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def set_section_geometry(section) -> None:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)


def configure_styles(doc: Document, compact=False) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
    normal.font.size = Pt(9.2 if compact else 9.7)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in [
        ("Title", 31, WHITE, 0, 8),
        ("Subtitle", 13, "D7C9FF", 0, 8),
        ("Heading 1", 22, PURPLE_DARK, 10, 7),
        ("Heading 2", 14.5, INK, 8, 4),
        ("Heading 3", 11.5, PURPLE_DARK, 6, 2),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display" if name != "Heading 3" else "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        st = styles[style_name]
        st.font.name = "Aptos"
        st.font.size = Pt(9.4 if compact else 9.8)
        st.font.color.rgb = RGBColor.from_string(INK)
        st.paragraph_format.space_after = Pt(3)


def configure_header_footer(doc: Document, short_title: str) -> None:
    for section in doc.sections:
        set_section_geometry(section)
        section.different_first_page_header_footer = True
        hdr = section.header
        table = hdr.add_table(rows=1, cols=2, width=Inches(7.0))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_cell_width(table.cell(0, 0), Inches(5.8))
        set_cell_width(table.cell(0, 1), Inches(1.2))
        p = table.cell(0, 0).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(short_title.upper())
        r.bold = True
        r.font.size = Pt(7.5)
        r.font.color.rgb = RGBColor.from_string(MUTED)
        q = table.cell(0, 1).paragraphs[0]
        q.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rr = q.add_run("NEURAI")
        rr.bold = True
        rr.font.size = Pt(8)
        rr.font.color.rgb = RGBColor.from_string(PURPLE)
        for c in table.rows[0].cells:
            cell_border(c, bottom={"val": "single", "sz": "10", "color": PURPLE})
            set_cell_margins(c, top=0, bottom=60, start=0, end=0)
        f = section.footer.paragraphs[0]
        f.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f.paragraph_format.space_before = Pt(0)
        r = f.add_run("NeurAI Platform  ·  26 August 2026  ·  ")
        r.font.size = Pt(7.5)
        r.font.color.rgb = RGBColor.from_string(MUTED)
        add_page_field(f)


def add_label(doc: Document, text: str, color=PURPLE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    keep_with_next(p)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(color)
    r.font.letter_spacing = Pt(0.6) if hasattr(r.font, "letter_spacing") else None


def add_kicker(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(SLATE)


def add_body(doc: Document, text: str, bold_lead: str | None = None, rtl=False) -> None:
    p = doc.add_paragraph()
    if rtl:
        set_rtl(p)
    if bold_lead and text.startswith(bold_lead):
        p.add_run(bold_lead).bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)


def add_bullets(doc: Document, items: Iterable[str], style="List Bullet") -> None:
    for item in items:
        p = doc.add_paragraph(style=style)
        p.add_run(item)


def add_callout(doc: Document, title: str, body: str, fill=PURPLE_PALE, accent=PURPLE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.right_indent = Cm(0.55)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_together = True
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "22")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), accent)
    borders.append(left)
    p_pr.append(borders)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "120")
    spacing.set(qn("w:after"), "120")
    p_pr.append(spacing)
    r = p.add_run(title + "  ")
    r.bold = True
    dark_fill = fill in {PURPLE_DARK, INDIGO, INDIGO}
    r.font.color.rgb = RGBColor.from_string(WHITE if dark_fill else PURPLE_DARK)
    body_run = p.add_run(body)
    body_run.font.color.rgb = RGBColor.from_string("E7E0F4" if dark_fill else INK)


def add_matrix(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[float] | None = None, font=8.2) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    if widths is None:
        widths = [16.5 / len(headers)] * len(headers)
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        set_cell_width(c, Cm(widths[i]))
        shade(c, PURPLE_DARK)
        set_cell_margins(c, 110, 110, 110, 110)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(font)
        r.font.color.rgb = RGBColor.from_string(WHITE)
    set_repeat_header(table.rows[0])
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for i, value in enumerate(row):
            set_cell_width(cells[i], Cm(widths[i]))
            shade(cells[i], WHITE if ridx % 2 == 0 else "F7F4FB")
            set_cell_margins(cells[i], 105, 105, 105, 105)
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            r.font.size = Pt(font)
            r.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_flow(doc: Document, steps: Sequence[tuple[str, str]], accent=PURPLE) -> None:
    table = doc.add_table(rows=2, cols=len(steps) * 2 - 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    total = len(steps) * 2 - 1
    for i, (title, body) in enumerate(steps):
        col = i * 2
        cell = table.cell(0, col)
        shade(cell, PURPLE_DARK if i % 2 == 0 else accent)
        set_cell_margins(cell, 130, 100, 110, 100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(8.7)
        r.font.color.rgb = RGBColor.from_string(WHITE)
        cell2 = table.cell(1, col)
        shade(cell2, PURPLE_PALE)
        set_cell_margins(cell2, 100, 90, 100, 90)
        p2 = cell2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p2.add_run(body)
        rr.font.size = Pt(7.3)
        rr.font.color.rgb = RGBColor.from_string(SLATE)
        for edge in [cell, cell2]:
            cell_border(edge, top={"val": "single", "sz": "6", "color": LINE}, bottom={"val": "single", "sz": "6", "color": LINE}, left={"val": "single", "sz": "6", "color": LINE}, right={"val": "single", "sz": "6", "color": LINE})
        if i < len(steps) - 1:
            a = table.cell(0, col + 1)
            a.merge(table.cell(1, col + 1))
            a.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p3 = a.paragraphs[0]
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r3 = p3.add_run("→")
            r3.font.size = Pt(16)
            r3.font.color.rgb = RGBColor.from_string(accent)
    for j in range(total):
        width = Cm(2.8 if j % 2 == 0 else 0.5)
        set_cell_width(table.cell(0, j), width)
        set_cell_width(table.cell(1, j), width)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_three_planes(doc: Document) -> None:
    rows = [
        ("Experience plane", "Next.js web · Persian/English shell · responsive recorder · assistant surfaces", CYAN),
        ("Control plane", "Fastify API · BFF boundary · agent runtime · workers · policy veto · audit", PURPLE),
        ("Data plane", "Supabase Postgres/Auth/Storage · RLS · grants · pgmq · provenance", CORAL),
    ]
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, (name, detail, color) in enumerate(rows):
        left, right = table.rows[i].cells
        set_cell_width(left, Cm(4.0)); set_cell_width(right, Cm(12.2))
        shade(left, color); shade(right, "F7F4FB")
        for c in (left, right):
            set_cell_margins(c, 180, 170, 180, 170)
            cell_border(c, top={"val": "single", "sz": "8", "color": WHITE}, bottom={"val": "single", "sz": "8", "color": WHITE}, left={"val": "nil"}, right={"val": "nil"})
        p = left.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(name); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor.from_string(WHITE if color != CYAN else PURPLE_DARK)
        q = right.paragraphs[0]
        rr = q.add_run(detail); rr.font.size = Pt(9); rr.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_screenshot(doc: Document, filename: str, caption: str, width=Inches(6.8)) -> None:
    path = SCREEN / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(path), width=width)
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(7)
    r = c.runs[0]
    r.italic = True
    r.font.size = Pt(7.5)
    r.font.color.rgb = RGBColor.from_string(MUTED)


def add_screenshot_pair(doc: Document, left: tuple[str, str], right: tuple[str, str]) -> None:
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, (filename, caption) in enumerate((left, right)):
        c = table.cell(0, i)
        set_cell_width(c, Cm(8.15)); set_cell_margins(c, 50, 50, 50, 50)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(SCREEN / filename), width=Inches(3.08))
        cap = table.cell(1, i)
        set_cell_width(cap, Cm(8.15)); set_cell_margins(cap, 40, 80, 80, 80)
        q = cap.paragraphs[0]; q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = q.add_run(caption); r.font.size = Pt(7.2); r.italic = True; r.font.color.rgb = RGBColor.from_string(MUTED)
        for x in (c, cap):
            cell_border(x, top={"val": "nil"}, bottom={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"})
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def section_page(doc: Document, number: str, title: str, kicker: str) -> None:
    doc.add_page_break()
    add_label(doc, number)
    h = doc.add_heading(title, level=1)
    h.paragraph_format.space_before = Pt(0)
    add_kicker(doc, kicker)


def cover(doc: Document, title: str, subtitle: str, deck=False) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = table.cell(0, 0)
    shade(c, INDIGO)
    set_cell_margins(c, 420, 400, 380, 400)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if BRAND.exists():
        p.add_run().add_picture(str(BRAND), width=Inches(0.72))
    p.add_run("\n")
    r = p.add_run("NEURAI PLATFORM")
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string(PURPLE)
    p.add_run("\n")
    r = p.add_run(title)
    r.bold = True; r.font.size = Pt(29 if not deck else 26); r.font.color.rgb = RGBColor.from_string(WHITE)
    p.add_run("\n")
    r = p.add_run(subtitle)
    r.font.size = Pt(12); r.font.color.rgb = RGBColor.from_string("D8CFFF")
    if HERO.exists():
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.add_run().add_picture(str(HERO), width=Inches(7.15))
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = line.add_run("Repository-grounded final edition  ·  26 August 2026")
    r.font.size = Pt(8.5); r.font.color.rgb = RGBColor.from_string(MUTED)
    line.add_run("\n")
    r = line.add_run("Architecture, product, trust, speech intelligence, market position, and expansion thesis")
    r.font.size = Pt(8.5); r.font.color.rgb = RGBColor.from_string(SLATE)


def build_guide() -> Path:
    doc = Document()
    configure_styles(doc)
    configure_header_footer(doc, "Platform technical & product guide")
    cover(doc, "The AI-native operating surface for organizational memory", "A complete product and architecture guide — beginning with Echo")

    section_page(doc, "00 · HOW TO READ THIS", "A final view—without pretending the future is already here", "One document that separates product truth, binding architecture, current code, and expansion intent.")
    add_callout(doc, "Snapshot discipline", "This edition is grounded in docs/SPEC.md, locked ARCHITECTURE.md decisions M1–M40, every tracked document under docs/, repository-wide route/schema/dependency inventories, and focused inspection of the live seams. It reflects commit 5c22477 plus the current local UI worktree on 26 August 2026.")
    add_matrix(doc, ["Label", "Meaning", "How it is used here"], [
        ["Implemented", "A route, data shape, UI surface, migration, or runtime path exists in the codebase.", "Described as a present capability; external deployment health is not inferred."],
        ["Architecturally committed", "A locked M-decision defines the required shape.", "Described as the product contract even if a specific provider or rollout remains conditional."],
        ["Expansion path", "The platform brief or AI-native plan points beyond Echo.", "Presented as thesis/roadmap, never as shipped fact."],
    ], widths=[3.2, 6.0, 7.3])
    add_body(doc, "Source hierarchy. Product behavior conflicts resolve to SPEC.md. System-shape conflicts resolve to the locked architecture. Code demonstrates the current implementation; a UI control or draft note never silently overrules those contracts.", bold_lead="Source hierarchy.")
    add_body(doc, "Naming. NeurAI Platform is the platform and shared shell. Echo is the call-intelligence app inside it. The Android recorder is Echo Mobile.", bold_lead="Naming.")

    section_page(doc, "01 · THE THESIS", "AI should be present before it is prompted", "NeurAI Platform turns lived work into governed memory—and governed memory into useful action.")
    add_body(doc, "Most assistants begin after a person opens a chat and reconstructs context. NeurAI Platform begins where organizational truth is created: conversations, decisions, commitments, corrections, and follow-through. It captures that truth with provenance, makes it searchable across time, and lets an agent act only through the authority of the person who asked.")
    add_callout(doc, "The north star", "Presence + hands + signals + bounded autonomy. The assistant sees what the person is already allowed to see, can operate the same product controls the person can operate, notices meaningful events, and asks before consequences when policy requires it.", fill="EEF8FF", accent=CYAN)
    add_matrix(doc, ["Old software", "AI-native NeurAI Platform"], [
        ["The person navigates to a feature.", "The assistant is present in the working surface."],
        ["The person retypes context into a prompt.", "Context is attached from governed organizational memory."],
        ["Automation is a hidden workflow.", "Actions are visible tool calls with consent and audit."],
        ["Permission is checked at the screen.", "Authority is enforced at API, database role, row policy, grant, and tool policy."],
        ["Knowledge fragments across apps.", "Echo creates the first high-fidelity memory substrate for future apps."],
    ], widths=[7.0, 9.5], font=8.7)
    add_body(doc, "This is pioneering not because it adds a chatbot to a dashboard, but because the platform treats memory, authority, and action as one product system.")

    section_page(doc, "02 · WHY ECHO FIRST", "The wedge is a conversation; the asset is organizational memory", "Calls are frequent, information-dense, expensive to forget, and naturally connected to future action.")
    add_flow(doc, [
        ("Capture", "Browser or Echo Mobile records real work"),
        ("Truth", "Versioned transcript with speakers and timing"),
        ("Memory", "Searchable summaries with provenance"),
        ("Reason", "Agent compares related calls and evidence"),
        ("Act", "Visible, governed tools and follow-through"),
    ])
    add_bullets(doc, [
        "Immediate value: searchable transcripts, summaries, corrections, speaker identity, and record management solve a clear daily problem.",
        "Compounding value: every processed call makes the assistant more context-rich for the next call, decision, and workflow.",
        "Trust laboratory: recording consent, sensitive content, tenancy, deletion, provenance, and hallucination risk force the platform to solve governance early.",
        "Platform proof: Echo exercises authentication, BFF boundaries, queues, storage, ML, LLM tools, connectors, administration, and responsive UX in one coherent app.",
        "Expansion bridge: once conversations become durable memory, the same substrate can power project, knowledge, workflow, and service apps without rebuilding identity or trust." 
    ])
    add_callout(doc, "Strategic choice", "A horizontal assistant without a truth source is generic. A point meeting-notes tool without a platform is bounded. Echo gives NeurAI Platform both a decisive starting product and the memory layer from which a broader AI-native workspace can grow.")

    section_page(doc, "03 · THE PRODUCT EXPERIENCE", "One calm shell; many specialized capabilities", "The user experiences a single assistant platform, while Echo provides the first deep domain surface.")
    add_screenshot_pair(doc, ("hub-fa.png", "Persian-first platform hub: prompt, sources, modes, and the Echo app entry."), ("new-meeting-recording.png", "Echo capture: waveform, live transcript, agenda, notes, and visible recording controls."))
    add_matrix(doc, ["Surface", "What it does", "Why it matters"], [
        ["Assistant hub", "Ask across permitted conversations, attach sources, select agents/models, use web search when enabled.", "Makes organizational memory conversational without separating it from governance."],
        ["Echo capture", "Create a meeting, choose language/template, record, pause/resume, show live captions, buffer parts.", "Reduces loss at the moment truth is created."],
        ["Records", "Filter, bulk act, archive, restore, scope, inspect processing state.", "Turns recordings into a manageable organizational corpus."],
        ["Record page", "Audio, transcript, speakers, notes, versioned summaries, translation/edit controls, assistant pane.", "Keeps source truth and derived insight visibly connected."],
        ["Speakers directory", "Enroll voice samples, merge people, group/read speaker history.", "Makes recurring participants recognizable without pretending uncertain labels are facts."],
        ["Skills / agents / workflows", "Configure instructions, models, tools, reusable agent cards and workflows.", "Moves from one generic assistant to governed, repeatable expertise."],
        ["Connectors / gateway", "API keys, webhooks, connector accounts and secret storage.", "Lets memory safely enter the rest of the organization’s toolchain."],
        ["Management / platform console", "Members, roles, models, server health, roots, org bootstrap and audits.", "Makes a multi-tenant system operable instead of merely demoable."],
    ], widths=[3.5, 6.2, 6.8], font=7.8)

    section_page(doc, "04 · FOUR PARTS, THREE PLANES", "A small number of explicit boundaries", "The platform is deliberately divided where trust, scale, and failure modes differ.")
    add_three_planes(doc)
    add_body(doc, "The four deployable parts are web/, core/, ml/, and Supabase. They form three planes: experience, control, and data. This is not microservice theatre: the ML boundary exists because speech dependencies and stateless media processing are operationally different; the BFF exists because browser credentials and internal topology should not leak; the worker exists because durable jobs have different lifecycle semantics from synchronous API requests.")
    add_flow(doc, [
        ("Browser", "UI + encrypted cookies + IndexedDB audio buffer"),
        ("BFF", "Same-origin API, auth method boundary, SSE proxy"),
        ("Core", "Fastify API, agent runtime, policy, job orchestration"),
        ("ML", "Stateless speech facade; no product authority"),
        ("Supabase", "Postgres, Auth, Storage, RLS, grants, pgmq"),
    ], accent=CORAL)
    add_callout(doc, "Why this shape", "Each boundary names a kind of authority. The browser holds a session, not provider secrets. ML receives a short-lived signed media URL, not product credentials. Workers re-resolve the record owner. Agent tools use the agent database role. The database remains the final wall.")

    section_page(doc, "05 · WEB & BFF", "The experience plane is also a security boundary", "Next.js renders the Persian-first shell while its BFF keeps tokens and internal services off the public edge.")
    add_matrix(doc, ["Component", "Responsibility", "Why chosen / trade-off"], [
        ["Next.js App Router", "Localized pages, server/client composition, route-level loading and error boundaries.", "A cohesive React full-stack shell with strong deployment ergonomics; requires discipline around server/client boundaries."],
        ["BFF (Backend for Frontend)", "Same-origin /api routes validate browser-facing input, manage auth cookies, proxy Core and SSE.", "Reduces exposed topology and prevents direct browser possession of internal credentials. It adds one hop, paid back in control."],
        ["next-intl", "Persian and English message catalogs, localized routing.", "Makes locale a structural dimension rather than scattered string conditionals."],
        ["Tailwind + shared scaffold", "Consistent tokens, menus, panels, RTL-aware spacing and responsive shell.", "Fast iteration with one source of layout truth; rendered checks remain essential because CSS intent can differ from computed output."],
        ["IndexedDB capture buffer", "Keeps audio parts resilient before upload/finish.", "Browser crashes and network loss should not erase the meeting; durable local buffering is worth the extra state machine."],
        ["SSE", "Streams assistant text/tool lifecycle and live captions downstream.", "Works through the BFF and Vercel constraints; simpler than a browser-to-Core WebSocket for this topology."],
    ], widths=[3.5, 6.3, 6.7], font=8.0)
    add_screenshot_pair(doc, ("records-bulk.png", "Records management: filters, selection, batch actions, and status visibility."), ("record-page.png", "Record detail: transcript source, audio, summary, notes, and contextual assistant."))

    section_page(doc, "06 · CORE API & WORKER", "One control plane, two execution tempos", "Fast synchronous request handling and durable asynchronous processing share identity rules but not failure semantics.")
    add_matrix(doc, ["Unit", "What it owns", "Importance"], [
        ["Fastify API", "Calls, parts, auth, assistant sessions, search, skills, agents, workflows, connectors, member/admin/platform routes.", "A narrow, typed control surface that makes authorization and errors explicit."],
        ["Identity resolver", "JWT sub → active app user → org, role, status; worker payload → call owner.", "There is no useful database handle until a person is known."],
        ["Lifecycle repositories", "Call/part state ladders, retry, missing-part handling, duration recompute, version pointers.", "Makes progress and recovery durable, visible, and idempotent."],
        ["pgmq worker", "Consumes process, speaker-link, summary, webhook, and signal jobs with retry/dead-letter behavior.", "Meeting processing must survive request lifetimes and machine restarts."],
        ["Agent runtime", "One Pi-based runtime for assistant and summarizer, with different tools and prompts.", "Avoids two drifting intelligence stacks and preserves run-level audit/replay."],
        ["Gateway", "Scoped API keys, signed webhooks, replay tolerance, redirect/SSRF controls.", "Connects the platform outward without weakening the inbound trust model."],
    ], widths=[3.2, 7.0, 6.3], font=8.0)
    add_callout(doc, "The worker runs as the owner", "A background job does not become a superuser because no browser is present. The payload carries the owner identity; Core re-resolves it and performs product work through the same RLS-bound door.", fill="FFF1EF", accent=CORAL)

    section_page(doc, "07 · ECHO DATA LIFECYCLE", "The transcript is truth; everything else can be rebuilt", "The pipeline preserves recoverability, provenance, and honest degradation at every stage.")
    add_flow(doc, [
        ("Create", "Call row + language + summary intent"),
        ("Upload parts", "Signed upload URLs; byte cap; offsets"),
        ("Process", "Transcode · VAD · STT · diarization"),
        ("Link", "Roster rows, optional voice matching"),
        ("Summarize", "Versioned artifact + grounding pass"),
    ])
    add_matrix(doc, ["Principle", "Implementation consequence"], [
        ["Transcript is source of truth", "Transcript segments and audio provenance outlive summary versions; a summary is derived and rebuildable."],
        ["Idempotency keys the artifact", "The worker checks for stored transcript rows rather than trusting a possibly stale status flag."],
        ["Missing is a state", "A failed/missing part can be marked explicitly, allowing the rest of a call to settle without inventing continuity."],
        ["Derived artifacts are versioned", "Human summary edits create new versions; the current pointer advances without erasing history."],
        ["Delete is role-aware and recoverable", "Soft deletion hides immediately and allows a 30-day restore; hard purge is not an agent tool."],
        ["Provenance travels with output", "Model, lane, timing granularity, diarization source, warnings, editor, and run records explain what produced the artifact."],
    ], widths=[5.0, 11.5], font=8.5)

    section_page(doc, "08 · SPEECH INTELLIGENCE", "Persian accuracy is a pipeline, not a checkbox", "Echo separates capture, speech detection, transcription, speaker structure, and voice identity so each can degrade honestly.")
    add_matrix(doc, ["Stage", "Current design", "Why"], [
        ["Transcode / probe", "ffmpeg normalizes media; duration/channel facts are measured.", "Provider inputs become predictable and oversize media fails before spend."],
        ["VAD", "Local speech detection removes meaningful silence; no-speech is logged and the original still reaches STT.", "Lower spend and latency without silently trusting an all-negative detector."],
        ["STT lanes", "Soniox primary, OpenRouter fallback; Persian/English hints and org glossary bias.", "Cloud accuracy with explicit fallback and user language steering."],
        ["Channel logic", "True two-channel audio maps speakers by microphone; dual-mono is detected and downmixed.", "Prevents duplicated words, invented speakers, and double billing."],
        ["Diarization", "Local sherpa-onnx CPU path for mono when needed; provider or channel labels otherwise.", "Speech structure stays productless and optional while avoiding a mandatory GPU service."],
        ["Voice identity", "Enrollment embeddings and conservative matching link call speakers to directory people.", "Recurring participants become useful context without claiming uncertain matches as fact."],
        ["Live captions", "Core holds the provider WebSocket; browser uploads audio and receives caption events via an ephemeral owned/ticketed session.", "Provider keys never enter the browser; captions remain content and never logs."],
    ], widths=[3.0, 7.3, 6.2], font=7.9)
    add_screenshot(doc, "speakers-enrollment.png", "Voice enrollment and speaker directory: identity is cultivated as a governed asset, not guessed from a filename.", width=Inches(6.35))

    section_page(doc, "09 · THE TIMING LADDER", "When fidelity falls, honesty remains", "A transcript must stay seekable even if a provider cannot return word-level timestamps.")
    add_flow(doc, [
        ("Word", "Exact word start/end; click-a-word UI"),
        ("Line", "Segment timing; click-a-line UI"),
        ("Anchored span", "First-to-last speech span for timeless text"),
    ], accent=CYAN)
    add_body(doc, "ML returns timings on the processed file’s zero-based timeline. Core anchors them to each call part’s offset. If a provider returns prose without time, Echo synthesizes one honest speech-span segment rather than zeros or an unseekable void. Non-empty segments with end ≤ start are refused.")
    add_callout(doc, "Why it matters", "Silently losing time makes a correct transcript feel broken and destroys citation. The timing ladder turns provider variation into an explicit product mode instead of a hidden defect.", fill="EEF8FF", accent=CYAN)
    add_matrix(doc, ["Granularity", "What the user can do", "What provenance says"], [
        ["Word", "Seek and highlight at the word level.", "timestamps: word; has_word_timestamps: true"],
        ["Segment", "Seek to a sentence/line span.", "timestamps: segment; feature demotes honestly"],
        ["None → anchored speech span", "Seek to the coarse speech region and read text.", "timestamps: none; degraded: true; real non-zero span"],
    ], widths=[4.5, 6.3, 5.7], font=8.5)

    section_page(doc, "10 · ONE AGENT RUNTIME", "Assistant and summarizer share the same governed intelligence", "Different jobs receive different prompts and tools, but every run is identity-bound, recorded, budgeted, and replayable.")
    add_flow(doc, [
        ("Identity", "Active actor + org + role"),
        ("Skill / agent", "Instructions, model, allowed tools"),
        ("Policy", "Role gates + tool ceilings + refusal budget"),
        ("Tools", "RLS-scoped reads or proposals/client actions"),
        ("Run record", "Prompt, model, steps, tokens, outcome"),
    ], accent=PURPLE)
    add_matrix(doc, ["Agent capability", "Mechanism", "Guardrail"], [
        ["Read organizational memory", "search_transcripts, read_window, get_call, list_related_calls", "Identity wrapper + RLS; cited call/time evidence."],
        ["Propose content changes", "correct_transcript, edit_speaker_roster, replace_summary", "Proposal first; explicit human confirmation applies through the app role."],
        ["Operate the visible UI", "navigation, capture controls, record/member/directory actions", "Only tools advertised by the surface; consent flags follow the autonomy setting."],
        ["Search the web", "OpenRouter online variant when the user enables it", "Base model passes catalogue rules before transport suffixing."],
        ["Summarize a call", "Same runtime with summarizer skill, related-call tools, templates, and grounding verifier", "Runs as the call owner; an empty result or missing system floor fails loudly."],
    ], widths=[4.2, 7.0, 5.3], font=8.1)
    add_callout(doc, "The profound choice", "The agent does not have its own omniscient identity. It borrows the caller’s authority and never more. That makes AI a participant inside the security model—not an exception around it.")

    section_page(doc, "11 · PRESENCE, HANDS, SIGNALS, AUTONOMY", "From chat feature to operating surface", "The platform’s AI-native layers answer four questions: where is the assistant, what can it touch, what does it notice, and when may it act?")
    add_matrix(doc, ["Layer", "Meaning", "Current architectural expression"], [
        ["Presence", "Assistant available in the hub and contextual panels.", "Shared session system, call context chips, visible tool lifecycle, persistent conversation history."],
        ["Hands", "Tools that read memory or operate the product surface.", "Server domain tools, propose-only write tools, and client-executed UI/write tools."],
        ["Signals", "Events that can create a brief, digest, or rule-triggered follow-up.", "call.processed, weekly digest and agent-rule queue patterns; event identity preserved."],
        ["Autonomy", "Watch → Assist → Act controls consequence, not intelligence.", "Watch offers no client tools; Assist requests consent for writes; Act is reserved for approved write classes."],
    ], widths=[3.2, 6.2, 7.1], font=8.4)
    add_flow(doc, [
        ("Watch", "Observe and answer; no client actions"),
        ("Assist", "Suggest and ask before writes"),
        ("Act", "Perform pre-approved write classes; still audited"),
    ], accent=CORAL)
    add_body(doc, "Autonomy is a policy dimension, not a personality setting. The same model can be helpful at every rung; what changes is which effects are exposed and when a person must consent.")

    section_page(doc, "12 · TRUST IS THE PRODUCT", "Prompts guide; database walls enforce", "Sensitive conversation data requires security properties that remain true even when application code or a model is wrong.")
    add_matrix(doc, ["Layer", "Enforcement", "Failure it prevents"], [
        ["JWT verification", "Signed access token verified against trusted JWKS; sub becomes identity candidate.", "Forged or expired browser identity."],
        ["Active user resolution", "App user, org, role, status checked before work.", "Pending, disabled, suspended, or unknown identities becoming product actors."],
        ["Identity-bound DB factory", "SET LOCAL ROLE + set_config(actor_id) inside the same transaction.", "Pooled connections leaking authority or code querying without a caller."],
        ["RLS", "Every protected row is filtered by actor/org/ownership/share policies.", "Cross-organization and cross-owner access."],
        ["Role grants", "echo_app and narrower echo_agent grants; agent has no DELETE.", "A model or tool gaining authority through a generous query path."],
        ["Central tool veto", "Declared-tool filter, admin gates, total/blocked attempt budgets.", "Tool discovery drift, permission grinding, and unbounded loops."],
        ["Proposal + consent", "Consequential content changes are proposed before app-role application.", "A plausible but wrong model edit becoming truth."],
        ["Audit/provenance", "Agent runs, steps, versions, editors, webhook delivery records.", "Invisible AI behavior and irreproducible derived artifacts."],
    ], widths=[3.4, 7.0, 6.1], font=7.7)
    add_callout(doc, "Non-negotiable invariant", "No database access without user identity. The agent borrows the caller’s authority. RLS plus grants are the wall. The agent role has no DELETE. Content never enters logs.", fill="FFF1EF", accent=RED)

    section_page(doc, "13 · DATA MODEL & PROVENANCE", "A record of truth, derivatives, actors, and decisions", "The schema is handwritten SQL because the wall must be reviewable at the level where it is enforced.")
    add_matrix(doc, ["Domain cluster", "Key entities", "Why they exist"], [
        ["Tenancy & identity", "org, app_user, invitation, status history, platform_operator", "One-org v1 membership, explicit roles/status, bootstrap and root audit."],
        ["Conversation truth", "call, call_part, transcript_segment, call_speaker, person", "Durable audio/transcript timeline, speaker roster, directory identity."],
        ["Derived intelligence", "summary, skill, agent_run", "Versioned summaries, reusable instructions, replayable model/tool execution."],
        ["Assistant", "agent_session, agent_message, feedback, share", "Persistent conversations with deletion semantics that preserve audit."],
        ["Automation", "assistant_agent, workflow_template, agent_card, agent_rule", "Named expertise, reusable process, proactive signals and autonomy."],
        ["Integration", "api_key, webhook, delivery, connector_connection, connector_secret", "External entry/exit with scoped authority and secret isolation."],
        ["Governance", "admin_action, platform_audit, deletion_record, proposal_decision", "Human and system accountability across sensitive actions."],
    ], widths=[3.8, 6.1, 6.6], font=7.9)
    add_body(doc, "Drizzle is used for queries only where appropriate; schema and security are handwritten numbered migrations. RLS and grant changes ship with SQL tests because the critical contract lives in Postgres, not in TypeScript types.")

    section_page(doc, "14 · PERSIAN-FIRST BY CONSTRUCTION", "RTL is only the visible layer", "Language, search, speech, dates, digits, names, and layouts all need first-class Persian behavior.")
    add_matrix(doc, ["Dimension", "Design requirement", "Why it matters"], [
        ["Direction", "RTL layout, start/end spacing, mirrored navigation, scoped LTR for ids/timestamps.", "A translated LTR shell still feels foreign and breaks mixed-script details."],
        ["Normalization", "Normalize Persian/Arabic character variants at ingest and query.", "Search should not miss the same word because of ی/ي or ک/ك variants."],
        ["Digits & dates", "Persian digits and Jalali-capable presentation with user preference.", "Dates and numbers are product semantics, not decoration."],
        ["Speech", "fa/en hints, mixed-language handling, glossary bias for names and org vocabulary.", "Proper names and code-switching are where generic transcription fails visibly."],
        ["Typography", "Persian-capable fonts and line-height tuned for Arabic script.", "Readable text density and diacritics require different visual rhythm."],
        ["Testing", "Sweep the structurally less-viewed locale and verify rendered state.", "A complete fa.json can hide a broken en.json—and the inverse."],
    ], widths=[3.1, 7.4, 6.0], font=8.2)
    p = doc.add_paragraph()
    set_rtl(p)
    r = p.add_run("نمونهٔ تجربه: «خلاصهٔ جلسه را با منبع و زمان نشان بده.»")
    r.bold = True; r.font.name = "Tahoma"; r.font.size = Pt(11); r.font.color.rgb = RGBColor.from_string(PURPLE_DARK)
    add_screenshot(doc, "hub-en.png", "The same platform shell in English: locale is a first-class route and message-catalog dimension, not an afterthought.", width=Inches(6.25))

    section_page(doc, "15 · CONNECTORS & GATEWAY", "Memory becomes valuable when it can travel safely", "The gateway is the platform’s contract with external systems; connectors are governed accounts, not copied secrets.")
    add_matrix(doc, ["Mechanism", "How it works", "Trust property"], [
        ["Scoped API keys", "Admin-created credentials with scopes, prefix/hash storage, last-used and revoke lifecycle.", "The raw key is shown once and cannot be recovered from the database."],
        ["Webhooks", "Subscribed events, signed deliveries, timestamped HMAC, replay tolerance, retries and delivery history.", "Consumers can verify origin; redirects and targets are re-checked against SSRF rules."],
        ["Connector accounts", "Org-scoped connection metadata; secret material stored through a dedicated secret backend.", "Product rows can describe a connection without exposing provider tokens."],
        ["OAuth / PKCE", "Provider authorization can bind a person or org without sharing passwords.", "Authorization code interception and browser secret exposure are reduced."],
        ["Workflow tools", "Agents/workflows select from approved tools and connector capabilities.", "An integration broadens utility without broadening the caller’s authority."],
    ], widths=[3.4, 7.4, 5.7], font=8.2)
    add_callout(doc, "Adapter rule", "Providers spell absence, errors, and retries differently. The adapter decides what 'already absent' means and returns a typed state; callers must not infer provider semantics from a raw status code.")

    section_page(doc, "16 · DEPLOYMENT & OPERATIONS", "Managed edges, self-controlled compute, one observable system", "The deployment follows the risk boundary: public UI at the edge, controlled Core/ML compute, managed durable data.")
    add_matrix(doc, ["Layer", "Placement", "Reason"], [
        ["Web / BFF", "Vercel", "Global delivery, same-origin browser boundary, managed build/deploy path."],
        ["Core API + worker", "Hetzner", "Long-lived processes, queue consumers, provider relays, explicit service supervision."],
        ["ML", "Hetzner CPU service", "Local ONNX/VAD/ffmpeg dependencies near Core without placing product state in ML."],
        ["Private ingress", "Cloudflare Tunnel", "No public origin exposure for internal services; managed edge routing."],
        ["Data/Auth/Storage", "Managed Supabase", "Postgres/RLS, authentication, object storage, and operational durability in one data plane."],
        ["Models", "Soniox + OpenRouter", "Specialized speech and user-selectable LLM lanes behind server-held keys."],
    ], widths=[3.4, 4.4, 8.7], font=8.3)
    add_bullets(doc, [
        "Health checks resolve the specific callable they guard; module presence is not health.",
        "Production-runtime boot tests start each process and make it answer one request.",
        "Logs carry codes, identifiers, counts, and structured database fields—never transcript content or signed URLs.",
        "Live network/spend tests are opt-in, run at package acceptance, and repeated at release gates.",
        "Runbooks separate deployed health from local code presence; a repository route is not proof that the external environment is healthy."
    ])

    section_page(doc, "17 · SMALL PARTS, CLEAR LANGUAGE", "A glossary for the whole platform", "The platform is sophisticated; its explanation should not require a security or ML background.")
    glossary = [
        ("BFF", "Backend for Frontend. A server layer built specifically for the web UI. It holds secure cookies, reshapes requests, and proxies Core so the browser does not talk directly to internal services."),
        ("JWT", "JSON Web Token. A signed, time-limited identity claim. Echo verifies the signature and reads the subject (sub); it does not trust a user id supplied in a request body."),
        ("JWKS", "JSON Web Key Set. The trusted public keys used to verify JWT signatures and key rotation."),
        ("RLS", "Row-Level Security. PostgreSQL policies that decide which rows the current actor may see or change—even if an application query forgets a filter."),
        ("RBAC", "Role-Based Access Control. Owner, admin, and member define broad authority; RLS and ownership add record-level precision."),
        ("Database grants", "Permissions attached to database roles. They are the hard ceiling: echo_agent has no DELETE even if a prompt or tool tries."),
        ("SSE", "Server-Sent Events. A one-way streaming HTTP connection used for assistant text, tool lifecycle, and caption events."),
        ("WebSocket", "A two-way long-lived connection. Core uses one to the live STT provider; the browser does not receive the provider key."),
        ("PKCE", "Proof Key for Code Exchange. An OAuth protection that ties the callback code to the browser session that started authorization."),
        ("HMAC", "A keyed signature over webhook content and time. Receivers can verify a delivery came from Echo and was not replayed outside the tolerance window."),
        ("SSRF", "Server-Side Request Forgery. An attack that tricks a server into calling internal/private addresses; webhook targets and redirects are checked to prevent it."),
        ("STT", "Speech-to-Text. The model/provider lane that turns audio into words, timing, language and sometimes speaker labels."),
        ("VAD", "Voice Activity Detection. A local model that identifies speech regions so long silence can be removed before paid transcription."),
        ("Diarization", "Answering ‘who spoke when?’ by clustering one audio stream into speaker labels such as S1 and S2. It is not the same as knowing the person’s name."),
        ("Voiceprint", "An embedding derived from voice samples and compared conservatively to recurring speakers. It supports a suggested match, not infallible identity."),
        ("Embedding", "A numeric vector representing features of speech or text. Similar vectors can indicate related voices or meaning."),
        ("LLM", "Large Language Model. Generates or interprets language; in NeurAI Platform it is surrounded by tools, identity, policy, provenance, and failure checks."),
        ("Tool calling", "The model asks software to perform a named operation with structured parameters. The platform still decides whether that tool is allowed and how it is executed."),
        ("Prompt injection", "Instructions hidden in retrieved content that try to redirect an agent. Echo quotes transcript content as data and enforces permission outside the prompt."),
        ("Provenance", "The record of what produced an artifact: provider/model, lane, timing mode, run, editor, warnings, and source."),
        ("Idempotency", "A retry produces one durable outcome, not duplicates. The worker checks the artifact that would prove prior success."),
        ("pgmq", "A durable Postgres-backed message queue used for background processing and retries."),
        ("DLQ", "Dead-letter queue. Jobs that exhausted safe retries land here for diagnosis instead of disappearing or looping forever."),
        ("Signed URL", "A short-lived capability to upload or download one object. ML can read the media it needs without holding platform storage credentials."),
        ("Soft delete", "Hide a record immediately but retain it for a recovery window; Echo uses a 30-day restore period before purge."),
        ("Tombstone", "A minimal durable record that an object was deleted, useful for audit and synchronization without retaining the original content."),
        ("Bidi / RTL", "Bidirectional and right-to-left layout behavior. Persian UI must place content and controls by start/end semantics while keeping ids and timestamps readable."),
        ("Jalali calendar", "The Persian calendar system supported as a display preference; storage remains unambiguous machine time."),
    ]
    add_matrix(doc, ["Term", "Plain-language explanation"], glossary, widths=[3.5, 13.0], font=7.9)

    section_page(doc, "18 · COMPETITIVE LANDSCAPE", "The category is proven; the opening is still distinct", "Gong, Otter, and Fireflies validate demand for conversation intelligence—NeurAI Platform competes on Persian-first depth, authority-aware action, and a broader platform thesis.")
    add_matrix(doc, ["Platform", "Best-known strength", "Closest overlap", "NeurAI Platform’s strategic contrast"], [
        ["Gong", "Revenue-focused conversation intelligence, deal signals, coaching, forecasting, embedded revenue agents.", "Capture, transcription, cross-conversation insight, action from interaction data.", "NeurAI is not limited to the revenue operating model: Echo is the first memory app inside a general AI-native platform, with caller-bound authority as a core primitive."],
        ["Otter.ai", "General meeting capture, live transcription, summaries, AI chat, enterprise workspace administration.", "Meeting knowledge, searchable transcripts, assistant Q&A, multi-device capture.", "Otter’s official May 2026 language list does not include Persian; NeurAI is Persian-first across UI, search, speech hints, dates, and organizational vocabulary."],
        ["Fireflies.ai", "Broad meeting automation, AskFred Q&A, AI skills, conversation intelligence, extensive integrations.", "Custom summaries/skills, cross-meeting questions, workflow/integration ecosystem.", "NeurAI makes product actions and data access flow through the same identity/RLS/grant wall, and treats autonomy as a visible Watch→Assist→Act policy."],
    ], widths=[2.7, 4.6, 4.6, 4.6], font=7.4)
    add_callout(doc, "How to compete", "Do not try to out-Gong Gong on revenue dashboards or out-integrate Fireflies on day one. Win the underserved Persian-first trust surface, make the agent visibly useful inside the product, and let Echo’s governed memory become the substrate for multiple apps.", fill="EEF8FF", accent=CYAN)
    add_body(doc, "This comparison reflects public positioning as of 26 August 2026. It is a strategic comparison, not an independent benchmark of accuracy, security certifications, or total feature coverage.")
    p = doc.add_paragraph()
    p.add_run("Official competitor sources: ").bold = True
    add_hyperlink(p, "Gong Conversation Intelligence", "https://www.gong.io/conversation-intelligence")
    p.add_run(" · ")
    add_hyperlink(p, "Otter home / Conversational Knowledge Engine", "https://otter.ai/")
    p.add_run(" · ")
    add_hyperlink(p, "Otter supported transcription languages", "https://help.otter.ai/hc/en-us/articles/26660468516631-Transcribe-conversations-in-English-Spanish-French-German-Japanese-or-Chinese-Simplified")
    p.add_run(" · ")
    add_hyperlink(p, "Fireflies platform", "https://fireflies.ai/")
    p.add_run(" · ")
    add_hyperlink(p, "Fireflies AskFred API", "https://docs.fireflies.ai/askfred/overview")

    section_page(doc, "19 · CHOICE RECORD", "Why these technologies—and what was consciously rejected", "The architecture is strongest where it records the trade-off instead of hiding it.")
    add_matrix(doc, ["Choice", "Selected", "Why", "Rejected / deferred"], [
        ["Web architecture", "Next.js + BFF", "One localized shell and secure same-origin boundary.", "Browser → Core direct calls; exposes internal topology and complicates auth."],
        ["API", "Fastify + Zod", "Typed, small, fast control surface with explicit validation.", "Framework-heavy service mesh before the domain needs it."],
        ["Database", "Supabase Postgres + handwritten SQL", "RLS/grants/schema visible and testable where enforced.", "ORM-owned schema that obscures the security wall."],
        ["Jobs", "pgmq", "Durable queue near transactional data with operational simplicity.", "In-memory jobs; they disappear at restart. Separate broker can come when scale proves it."],
        ["Agent runtime", "Pi + central policy + typed tools", "One runtime for assistant/summarizer with tool-call interception and replay.", "Prompt-only permission; two independent agent stacks."],
        ["Speech", "Soniox + OpenRouter fallback + local ONNX", "Persian-capable cloud STT with stateless local speech structure.", "Browser provider keys; a stateful product-aware ML service."],
        ["Live transport", "Chunk upload + SSE / ephemeral ticket", "Works with BFF/Vercel topology while Core owns provider socket.", "Browser-to-provider WebSocket containing provider credential."],
        ["Delete", "Soft delete + named purge doors", "Recoverability, role hierarchy, audit, and no agent DELETE.", "Generic DELETE permission or hidden service-account bypass."],
    ], widths=[2.8, 3.6, 5.6, 4.5], font=7.6)

    section_page(doc, "20 · STATUS & EXPANSION", "Build the substrate once; compound it through apps", "Echo proves the hard platform capabilities. Future apps should reuse identity, memory, tools, connectors, and governance—not fork them.")
    add_matrix(doc, ["Now: Echo + shared platform", "Next: deepen", "Then: expand"], [
        ["Capture, transcript, speaker structure, versioned summaries, search, assistant sessions, skills/agents/workflows, connectors/gateway, management and platform-root surfaces.", "Acceptance-grade Persian speech quality, stronger templates/grounding, operational SLOs, richer signals, approved Act classes, production connector catalogue.", "Project memory, knowledge/wiki, service workflows, org briefs, cross-app actions—all on the same identity and memory substrate."],
    ], widths=[5.5, 5.5, 5.5], font=8.2)
    add_flow(doc, [
        ("Echo", "Conversation truth and durable memory"),
        ("Shared agent", "Cross-memory reasoning and product tools"),
        ("Signals", "Proactive briefs and follow-through"),
        ("More apps", "Projects, knowledge, service, operations"),
    ])
    add_callout(doc, "Expansion rule", "A new app is valuable when it contributes a new source of truth or a new governed action surface. It should not create a second identity system, a second memory store, or an unbounded agent.")

    section_page(doc, "21 · CLOSING", "Build the AI people can trust with the work that matters", "NeurAI Platform starts with the most human data in an organization—conversation—and turns it into memory without giving up authority, provenance, or language.")
    add_body(doc, "Echo is the first app because it transforms an urgent loss—forgotten conversations—into a compounding asset. NeurAI Platform is the larger bet: when every memory is attributable, every action is bounded, and every interface is built for the person’s language, AI stops being a side chat and becomes a trustworthy operating surface.")
    add_callout(doc, "The promise", "Capture what happened. Explain what it means. Act only with authority. Learn across time. Do it naturally in Persian. Then let every future app inherit that trust.", fill=PURPLE_DARK, accent=CORAL)
    add_body(doc, "That is why the platform begins with Echo—and why Echo is much more than meeting notes.")

    section_page(doc, "APPENDIX · SOURCE REGISTER", "What this edition is grounded in", "Primary project sources first; official competitor sources only for market positioning.")
    add_matrix(doc, ["Source", "Role"], [
        ["docs/SPEC.md", "Product behavior source of truth."],
        ["ARCHITECTURE.md", "Locked implementation decisions M1–M40 and amendments."],
        ["docs/PLATFORM-BRIEF.md", "Platform narrative, app model, and strategic framing."],
        ["docs/AI-NATIVE-PLAN.md", "Presence, hands, signals, autonomy and sequence."],
        ["docs/CONNECTORS.md", "Connector boundary and product vocabulary."],
        ["docs/PLATFORM-ROOT.md", "Bootstrap/root administration behavior."],
        ["docs/PLATFORM-OPERATIONS-RUNBOOK.md", "Deployment and operational practices."],
        ["docs/CLOSE-m4-frontend.md", "Frontend closure evidence and UI implementation notes."],
        ["docs/screenshots/*", "Current visual evidence of the web product surfaces."],
        ["web/, core/, ml/, db/", "Route, dependency, runtime, schema, security and current capability evidence."],
        ["Gong / Otter / Fireflies official product and help pages", "Competitive public positioning as of 26 August 2026."],
    ], widths=[6.5, 10.0], font=8.4)
    add_body(doc, "Important: repository evidence is not a live production probe. Where the code contains a capability but external provider credentials, migrations, models, or deployed versions determine availability, this document describes the architecture/current implementation rather than claiming a verified customer-facing service level.")

    path = OUT / "NeurAI-Platform-Technical-and-Product-Guide.docx"
    doc.save(path)
    return path


SLIDES = [
    {
        "n": 1,
        "title": "NeurAI Platform",
        "claim": "The AI-native operating surface for organizational memory",
        "time": "0:00–1:10",
        "goal": "Open with the platform ambition, then make trust and Persian-first design part of the promise—not footnotes.",
        "speech": [
            "What if AI did not begin when we opened a chat? What if it was already present where the work happened—listening with consent, remembering with provenance, and acting only with our authority? That is NeurAI Platform.",
            "We are building an AI-native operating surface for organizations. It captures the truth of work, makes that truth usable across time, and gives people an assistant that can reason and act without becoming a security exception.",
            "Our first app is Echo. Echo turns calls and meetings into transcripts, speakers, versioned summaries, searchable memory, and governed actions. Persian is not a translated skin here. The interface, speech pipeline, search behavior, names, dates, digits, and direction all begin Persian-first.",
            "The promise is simple: capture what happened, explain what it means, act only with authority, and learn across time."],
        "cue": "Pause after ‘That is NeurAI Platform.’ Let the cover image establish scale before naming Echo.",
        "transition": "To understand why this matters, start with the loss every organization already accepts as normal."},
    {
        "n": 2,
        "title": "The problem is not more meetings—it is memory loss",
        "claim": "Truth is created in conversation and disappears into fragments",
        "time": "1:10–2:20",
        "goal": "Turn meeting pain into a strategic data problem.",
        "speech": [
            "Organizations do not suffer because they lack information. They suffer because the most important information is born in conversation and then decays.",
            "A commitment lives in someone’s notes. A customer objection stays in a recording. A correction is repeated three meetings later. Context is retyped into a generic assistant, usually without source, scope, or history.",
            "Traditional meeting tools stop at notes. Traditional enterprise systems start after the human has entered structured data. The gap between those two moments is where decisions, trust, and time are lost.",
            "NeurAI Platform treats this as a memory architecture problem. We need a source of truth, durable provenance, identity-aware access, and a path from evidence to action."],
        "cue": "Point to the four losses: source, ownership, continuity, follow-through.",
        "transition": "That leads to the most important product decision we made: where to begin."},
    {
        "n": 3,
        "title": "Echo is the wedge—and the memory substrate",
        "claim": "One conversation can become truth, memory, reasoning, and action",
        "time": "2:20–3:35",
        "goal": "Explain the why-now and why-Echo strategy.",
        "speech": [
            "We start with Echo because conversations are high-frequency, information-dense, and expensive to forget. The value is immediate: record, transcribe, identify speakers, correct, summarize, search, and share.",
            "But the strategic value compounds. The transcript becomes durable memory. The next call can be understood against previous calls. The assistant can cite the moment a decision was made. A workflow can react to a processed call without scraping a dashboard.",
            "Echo also forces us to solve the hard platform problems early: recording consent, sensitive content, multi-tenant isolation, background processing, model provenance, deletion, human correction, and bounded action.",
            "A horizontal assistant without a truth source is generic. A meeting tool without a platform is bounded. Echo gives us both the wedge and the substrate."],
        "cue": "Walk left to right through Capture → Truth → Memory → Reason → Act.",
        "transition": "Now let’s make that concrete in the experience a user sees."},
    {
        "n": 4,
        "title": "One calm product surface",
        "claim": "The platform hub and Echo feel like one system—not a bundle of AI features",
        "time": "3:35–5:00",
        "goal": "Demo the current experience through two screenshots.",
        "speech": [
            "This is the Persian-first platform hub. The user can ask across permitted memory, attach calls as sources, select an agent or model, and choose whether web search is part of the turn.",
            "Then Echo moves to the moment of truth creation. The recorder keeps waveform, live transcript, agenda, and notes in one place. Audio is buffered locally in the browser, uploaded in bounded parts, and processed without exposing speech-provider credentials.",
            "After processing, the record page keeps audio, transcript, speaker roster, notes, summary versions, and the contextual assistant together. The user never has to choose between evidence and intelligence.",
            "The visual language is intentionally calm: dark neutral surfaces, precise violet emphasis, strong hierarchy, and RTL behavior that feels native rather than mirrored after the fact."],
        "cue": "Show the hub first; then the recorder. If live product is available, use the screenshots as the fallback path.",
        "transition": "The visible product is calm because the intelligence underneath has a very explicit shape."},
    {
        "n": 5,
        "title": "AI-native means presence + hands + signals",
        "claim": "The assistant moves from chat to a governed operating layer",
        "time": "5:00–6:15",
        "goal": "Explain the platform’s AI-native primitives in memorable language.",
        "speech": [
            "Our definition of AI-native has three active layers and one control.",
            "Presence means the assistant is available in the hub and beside the record where context already exists. Hands means it can use tools: search transcripts, read related calls, navigate the product, manage records, or propose a correction. Signals mean important events can create a brief or trigger a governed rule without waiting for a fresh prompt.",
            "The control is autonomy: Watch, Assist, Act. Watch observes and answers. Assist proposes and asks before writes. Act is reserved for pre-approved write classes and remains fully audited.",
            "This is a crucial distinction: autonomy changes consequences, not intelligence. We do not make the model ‘more trusted.’ We make the exposed effects more explicit."],
        "cue": "Emphasize the words Presence, Hands, Signals, Autonomy. They should become the audience’s vocabulary.",
        "transition": "Those primitives work because the architecture gives every kind of authority a place."},
    {
        "n": 6,
        "title": "Four parts. Three planes. Explicit boundaries.",
        "claim": "Experience, control, and data scale independently without losing identity",
        "time": "6:15–7:50",
        "goal": "Give technical confidence without drowning the room in boxes.",
        "speech": [
            "The platform has four deployable parts. The web application is the experience plane. Its Backend for Frontend keeps secure cookies and internal topology away from the browser. Core is the control plane: the Fastify API, worker, agent runtime, policies, gateway, and live relay. ML is a stateless speech facade. Supabase is the durable data plane: Postgres, Auth, Storage, row-level security, grants, and queues.",
            "The boundaries are about authority. The browser holds a session, not provider secrets. ML gets a short-lived signed URL to one audio object, not a platform credential. Background jobs re-resolve the call owner. Agent tools use a narrower database role.",
            "We chose a small number of real boundaries instead of a large number of nominal services. That keeps the system operable while making the trust seams reviewable."],
        "cue": "Trace one request: Browser → BFF → Core → identity-bound database. Then trace one audio job through Worker → signed URL → ML.",
        "transition": "Now we can show the differentiator that is hardest to copy with a feature checklist: trust as architecture."},
    {
        "n": 7,
        "title": "The agent borrows authority—never invents it",
        "claim": "JWT, identity, RLS, grants, policy, and consent form one wall",
        "time": "7:50–9:10",
        "goal": "Make the security model understandable and distinctive.",
        "speech": [
            "A JWT is the signed token that tells us who the caller claims to be. We verify it, resolve an active platform user, and attach that identity inside the same database transaction as every query.",
            "Postgres row-level security limits which rows that person can see. Role grants set the ceiling on what the application and agent can do. The agent role has no DELETE. A central tool policy checks declared tools, admin gates, and attempt budgets. Consequential content changes become proposals that a person confirms.",
            "No single layer carries the promise. If UI code is wrong, RLS still applies. If a prompt is manipulated, grants still apply. If a model keeps trying, the tool budget ends the loop.",
            "The profound choice is that the agent has no omniscient service identity. It borrows the caller’s authority and never more."],
        "cue": "Build the wall bottom-up. End on ‘never more’ and pause.",
        "transition": "That same honesty appears in the most technically demanding part of Echo: Persian speech."},
    {
        "n": 8,
        "title": "Persian-first speech—with honest degradation",
        "claim": "Accuracy is a pipeline; trust survives when fidelity changes",
        "time": "9:10–10:35",
        "goal": "Show product and engineering depth around speech.",
        "speech": [
            "Persian-first is not a language toggle. It shapes the UI, text normalization, digits, dates, mixed-script typography, search, speech hints, and the organization’s glossary of names and terms.",
            "The audio pipeline measures channels, avoids the dual-mono trap, runs local voice activity detection, uses Soniox as the primary STT lane with an OpenRouter fallback, and adds local diarization where a mono stream needs speaker structure. Voice enrollment can conservatively connect recurring speakers to directory people.",
            "Most importantly, fidelity degrades honestly. Word timestamps give click-a-word. Segment timestamps give click-a-line. If a provider gives prose with no timing, Echo anchors it to a real speech span instead of returning zeros or nothing.",
            "A correct transcript that cannot be traced back to audio is not complete. Our timing ladder keeps evidence usable."],
        "cue": "Use the timing ladder as the memorable technical example: word → line → anchored span.",
        "transition": "With that product and architecture in mind, how do we position against a proven category?"},
    {
        "n": 9,
        "title": "The category is proven; our opening is distinct",
        "claim": "Gong, Otter, and Fireflies validate demand—NeurAI changes the axis",
        "time": "10:35–12:05",
        "goal": "Frame competition with respect and a specific wedge.",
        "speech": [
            "Gong proves that conversation intelligence can become an operating system for revenue. Otter proves that general meeting knowledge can become conversational. Fireflies proves the pull for integrations, skills, and meeting automation.",
            "We should respect those strengths. Our opportunity is not to copy every checklist. NeurAI Platform is differentiated by three connected choices: Persian-first depth, authority-aware agents, and a platform designed to grow beyond meeting notes.",
            "Otter’s current official language list does not include Persian. Fireflies and Gong are not publicly positioned as Persian-first operating systems. More importantly, our agent, worker, and database all carry the same caller-bound authority model.",
            "Our strategy is to win an underserved trust-and-language surface, then compound Echo’s memory into more apps and actions."],
        "cue": "Do not disparage competitors. Acknowledge strengths, then name our three axes.",
        "transition": "The final slide is the expansion thesis: what becomes possible once Echo works."},
    {
        "n": 10,
        "title": "Echo is the beginning, not the boundary",
        "claim": "One governed memory substrate can power an entire AI-native platform",
        "time": "12:05–13:20",
        "goal": "Close on inevitability and disciplined ambition.",
        "speech": [
            "Echo begins with conversation truth. The shared agent turns that truth into reasoning. Signals turn reasoning into timely follow-through. Then future apps—projects, knowledge, service, operations—can reuse the same identity, memory, tools, connectors, and audit model.",
            "That reuse is the platform advantage. A new app should contribute a new source of truth or a new governed action surface. It should not create a second identity system, a second memory store, or an unbounded agent.",
            "We are building the AI people can trust with the work that matters: present before the prompt, natural in Persian, evidence-backed, and powerful without becoming omnipotent.",
            "Capture what happened. Explain what it means. Act only with authority. Learn across time. That is NeurAI Platform—and Echo is how it begins."],
        "cue": "Finish slowly. Leave the four-sentence promise on screen for questions.",
        "transition": "Invite questions on the wedge, the trust model, or the expansion sequence."},
]


def build_speaker_script() -> Path:
    doc = Document()
    configure_styles(doc, compact=True)
    configure_header_footer(doc, "10-slide demo speaker script")
    cover(doc, "Executive demo — speaker script", "A dedicated, timed narrative for each of the 10 presentation slides", deck=True)
    doc.add_page_break()
    add_label(doc, "PRESENTATION MAP")
    doc.add_heading("A 13-minute story with room for a live product moment", level=1)
    add_kicker(doc, "Problem → wedge → product → AI-native model → architecture → trust → Persian speech → competition → expansion")
    add_matrix(doc, ["Slide", "Claim", "Time"], [[str(s["n"]), s["claim"], s["time"]] for s in SLIDES], widths=[1.4, 12.8, 2.3], font=8.1)
    add_callout(doc, "Delivery style", "Excited, precise, and credible. Treat every diagram as a story, not a checklist. Keep the strongest words—memory, authority, Persian-first, presence—consistent across the talk.")
    for s in SLIDES:
        doc.add_page_break()
        add_label(doc, f"SLIDE {s['n']}  ·  {s['time']}")
        doc.add_heading(s["title"], level=1)
        add_kicker(doc, s["claim"])
        add_callout(doc, "Purpose", s["goal"], fill="EEF8FF", accent=CYAN)
        doc.add_heading("Suggested speech", level=2)
        for para in s["speech"]:
            add_body(doc, para)
        doc.add_heading("Presenter cue", level=2)
        add_body(doc, s["cue"])
        doc.add_heading("Transition", level=2)
        p = doc.add_paragraph()
        r = p.add_run(s["transition"])
        r.italic = True
        r.font.color.rgb = RGBColor.from_string(PURPLE_DARK)
    doc.add_page_break()
    add_label(doc, "Q&A BACK POCKET")
    doc.add_heading("Short answers to likely questions", level=1)
    add_matrix(doc, ["Question", "Answer spine"], [
        ["Why not begin with a generic assistant?", "Without a governed truth source, the assistant is generic and context has to be re-entered. Echo creates the first compounding memory asset."],
        ["Why not use one vendor for all AI?", "Speech and reasoning have different quality/cost/failure profiles. The adapters keep providers replaceable while provenance makes the chosen lane visible."],
        ["Can the agent delete data?", "No. The agent database role has no DELETE. Soft-delete UI actions are surface tools under the caller’s authority; permanent purge is a named administrative path."],
        ["What makes Persian-first defensible?", "It spans capture, UI direction, normalization, search, mixed-language transcription, names/glossary, digits/dates, typography, and acceptance fixtures—not a translation file."],
        ["What is shipped versus planned?", "The repository contains the Echo and shared platform surfaces, routes, schema, queues, agent runtime, connectors and admin paths. Expansion apps are a disciplined thesis, not claimed as shipped."],
        ["Why Supabase/Postgres?", "Identity, durable data, storage, queueing, RLS and grants can form one reviewable data plane. The database can enforce the wall even when application code is wrong."],
    ], widths=[5.2, 11.3], font=8.1)
    path = OUT / "NeurAI-Platform-Demo-Speaker-Script.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    print(build_guide())
    print(build_speaker_script())
