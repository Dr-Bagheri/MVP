from pathlib import Path
import re
import zipfile

from docx import Document
from pypdf import PdfReader


ROOT = Path(r"C:\Users\amirreza\Desktop\mvp\deliverables\neurai-platform-final")
GUIDE = ROOT / "NeurAI-Platform-Technical-and-Product-Guide.docx"
DECK = ROOT / "NeurAI-Platform-Executive-Demo-10-Slides.pptx"
SCRIPT = ROOT / "NeurAI-Platform-Demo-Speaker-Script.docx"


for artifact in (GUIDE, DECK, SCRIPT):
    with zipfile.ZipFile(artifact) as archive:
        print(f"zip_integrity {artifact.name}: {archive.testzip() or 'ok'}")

for artifact in (GUIDE, SCRIPT):
    document = Document(artifact)
    text = "\n".join(p.text for p in document.paragraphs)
    print(
        f"docx_blocks {artifact.name}: "
        f"paragraphs={len(document.paragraphs)} "
        f"tables={len(document.tables)} images={len(document.inline_shapes)}"
    )
    hits = [needle for needle in ("TODO", "TBD", "Lorem ipsum", "PLACEHOLDER") if needle in text]
    print(f"placeholder_hits {artifact.name}: {hits or 'none'}")

with zipfile.ZipFile(DECK) as archive:
    names = archive.namelist()
    slides = [name for name in names if re.fullmatch(r"ppt/slides/slide\d+\.xml", name)]
    notes = [name for name in names if re.fullmatch(r"ppt/notesSlides/notesSlide\d+\.xml", name)]
    xml_text = "".join(
        archive.read(name).decode("utf-8", errors="ignore")
        for name in slides + notes
    )
    hits = [needle for needle in ("TODO", "TBD", "Lorem ipsum", "PLACEHOLDER") if needle in xml_text]
    print(f"pptx_structure slides={len(slides)} notes={len(notes)}")
    print(f"placeholder_hits {DECK.name}: {hits or 'none'}")

guide_pdf = ROOT / "build" / "qa-guide" / "NeurAI-Platform-Technical-and-Product-Guide.pdf"
script_pdf = ROOT / "build" / "qa-script" / "NeurAI-Platform-Demo-Speaker-Script.pdf"
print(f"pdf_pages guide={len(PdfReader(guide_pdf).pages)} script={len(PdfReader(script_pdf).pages)}")
